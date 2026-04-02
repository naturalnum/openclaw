#!/usr/bin/env python3

from __future__ import annotations

import argparse
from copy import deepcopy
import json
import re
import subprocess
import tempfile
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.text import WD_TAB_LEADER
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Mm, Pt
from pypdf import PdfReader

SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_ROOT = SCRIPT_DIR.parent
DEFAULT_TEMPLATE = SKILL_ROOT / "references" / "纵向科技项目技术报告编写模板-2023版（新模板）.docx"
DEFAULT_FRONT_REFERENCE = SKILL_ROOT / "references" / "front-matter-fixed.docx"


def normalize(text: str) -> str:
    return " ".join(text.split())


def clear_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def clear_paragraph_content(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)


def first_nonempty_run(paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            return run
    return paragraph.runs[0] if paragraph.runs else None


def set_rfonts(target, ascii_font=None, hansi_font=None, east_asia_font=None, cs_font=None) -> None:
    if hasattr(target, "_r"):
        rPr = target._r.get_or_add_rPr()
    else:
        rPr = target.element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    if ascii_font:
        rFonts.set(qn("w:ascii"), ascii_font)
    if hansi_font:
        rFonts.set(qn("w:hAnsi"), hansi_font)
    if east_asia_font:
        rFonts.set(qn("w:eastAsia"), east_asia_font)
    if cs_font:
        rFonts.set(qn("w:cs"), cs_font)


def copy_run_font(dst_run, src_run, east_asia_font=None) -> None:
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    dst_run.font.name = src_run.font.name
    dst_run.font.size = src_run.font.size
    src_rpr = src_run._r.rPr
    if src_rpr is not None and src_rpr.rFonts is not None:
        set_rfonts(
            dst_run,
            ascii_font=src_rpr.rFonts.get(qn("w:ascii")),
            hansi_font=src_rpr.rFonts.get(qn("w:hAnsi")),
            east_asia_font=east_asia_font or src_rpr.rFonts.get(qn("w:eastAsia")),
            cs_font=src_rpr.rFonts.get(qn("w:cs")),
        )


def copy_paragraph_format(dst, src) -> None:
    dst.alignment = src.alignment
    dpf = dst.paragraph_format
    spf = src.paragraph_format
    dpf.left_indent = spf.left_indent
    dpf.right_indent = spf.right_indent
    dpf.first_line_indent = spf.first_line_indent
    dpf.space_before = spf.space_before
    dpf.space_after = spf.space_after
    dpf.line_spacing = spf.line_spacing
    dpf.line_spacing_rule = spf.line_spacing_rule


def ensure_style(doc: Document, style_name: str, sample_paragraph, east_asia_font: str) -> None:
    names = {style.name for style in doc.styles}
    if style_name in names:
        style = doc.styles[style_name]
    else:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles["Normal"]
    copy_paragraph_format(style, sample_paragraph)
    sample_run = first_nonempty_run(sample_paragraph)
    if sample_run is not None:
        style.font.name = sample_run.font.name
        style.font.size = sample_run.font.size
        style.font.bold = sample_run.bold
        style.font.italic = sample_run.italic
        src_rpr = sample_run._r.rPr
        if src_rpr is not None and src_rpr.rFonts is not None:
            set_rfonts(
                style,
                ascii_font=src_rpr.rFonts.get(qn("w:ascii")),
                hansi_font=src_rpr.rFonts.get(qn("w:hAnsi")),
                east_asia_font=east_asia_font,
                cs_font=src_rpr.rFonts.get(qn("w:cs")),
            )
    pPr = style.element.get_or_add_pPr()
    if pPr.numPr is not None:
        pPr.remove(pPr.numPr)


def ensure_toc_style(doc: Document, style_name: str, left_indent_pt: float = 0.0) -> None:
    names = {style.name for style in doc.styles}
    if style_name in names:
        style = doc.styles[style_name]
    else:
        style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.font.bold = False
    set_rfonts(style, ascii_font="宋体", hansi_font="宋体", east_asia_font="宋体", cs_font="宋体")
    p = style.paragraph_format
    p.first_line_indent = None
    p.left_indent = Pt(left_indent_pt) if left_indent_pt else None
    p.space_before = Pt(0)
    p.space_after = Pt(0)
    p.line_spacing = 1.5


def copy_style_format(dst_style, src_style) -> None:
    dst_el = dst_style.element
    src_el = src_style.element
    for tag in [qn("w:pPr"), qn("w:rPr")]:
        existing = dst_el.find(tag)
        if existing is not None:
            dst_el.remove(existing)
        src_child = src_el.find(tag)
        if src_child is not None:
            dst_el.append(deepcopy(src_child))


def force_zero_indent(style) -> None:
    pPr = style.element.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:right"), "0")
    ind.set(qn("w:firstLine"), "0")
    ind.set(qn("w:hanging"), "0")
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), "left")
    text_alignment = pPr.find(qn("w:textAlignment"))
    if text_alignment is not None:
        pPr.remove(text_alignment)


def remove_style_numbering(style) -> None:
    pPr = style.element.get_or_add_pPr()
    if pPr.numPr is not None:
        pPr.remove(pPr.numPr)


def ensure_report_styles(doc: Document, template: Document | None = None) -> None:
    h1 = doc.styles["Heading 1"]
    h2 = doc.styles["Heading 2"]
    h3 = doc.styles["Heading 3"]

    if template is not None:
        copy_style_format(h1, template.styles["Heading 1"])
        copy_style_format(h2, template.styles["Heading 2"])
        copy_style_format(h3, template.styles["3级"])
        h3.base_style = doc.styles["Normal"]

    h1.font.name = "黑体"
    h1.font.size = Pt(15)
    h1.font.bold = True
    set_rfonts(h1, ascii_font="黑体", hansi_font="黑体", east_asia_font="黑体", cs_font="黑体")
    h1pf = h1.paragraph_format
    h1pf.left_indent = Pt(0)
    h1pf.right_indent = Pt(0)
    h1pf.first_line_indent = Pt(0)
    force_zero_indent(h1)
    remove_style_numbering(h1)

    h2.font.name = "黑体"
    h2.font.size = Pt(14)
    h2.font.bold = True
    set_rfonts(h2, ascii_font="黑体", hansi_font="黑体", east_asia_font="黑体", cs_font="黑体")
    h2pf = h2.paragraph_format
    h2pf.left_indent = Pt(0)
    h2pf.right_indent = Pt(0)
    h2pf.first_line_indent = Pt(0)
    force_zero_indent(h2)
    remove_style_numbering(h2)

    h3.font.name = "黑体"
    h3.font.size = Pt(12)
    h3.font.bold = False
    set_rfonts(h3, ascii_font="黑体", hansi_font="黑体", east_asia_font="黑体", cs_font="黑体")
    h3pf = h3.paragraph_format
    h3pf.left_indent = Pt(0)
    h3pf.right_indent = Pt(0)
    h3pf.first_line_indent = Pt(0)
    force_zero_indent(h3)
    remove_style_numbering(h3)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal.font.bold = False
    set_rfonts(normal, ascii_font="宋体", hansi_font="宋体", east_asia_font="宋体", cs_font="宋体")
    npf = normal.paragraph_format
    npf.first_line_indent = Pt(24)
    npf.space_before = Pt(0)
    npf.space_after = Pt(0)
    npf.line_spacing = 1.5

    ensure_toc_style(doc, "TOC 1", 0)
    ensure_toc_style(doc, "TOC 2", 24)
    ensure_toc_style(doc, "TOC 3", 48)


def apply_document_layout(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Mm(30)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(25)
        section.right_margin = Mm(25)
        section.footer_distance = Mm(15)


def add_paragraph(doc: Document, text: str, style_name: str, sample_paragraph, east_asia_font: str, copy_format: bool = True):
    p = doc.add_paragraph(style=style_name)
    if copy_format:
        copy_paragraph_format(p, sample_paragraph)
    clear_runs(p)
    run = p.add_run(text)
    sample_run = first_nonempty_run(sample_paragraph)
    if sample_run is not None:
        copy_run_font(run, sample_run, east_asia_font=east_asia_font)
    return p


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_bookmark(paragraph, bookmark_id: int, name: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_section_page_number(section, fmt: str | None = None, start: int | None = None) -> None:
    sectPr = section._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sectPr.append(pg)
    if fmt:
        pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))


def get_section_page_number(section) -> tuple[str | None, int | None]:
    sect_pr = section._sectPr
    pg = sect_pr.find(qn("w:pgNumType"))
    if pg is None:
        return None, None
    fmt = pg.get(qn("w:fmt"))
    start = pg.get(qn("w:start"))
    return fmt, int(start) if start is not None else None


def add_page_field(paragraph) -> None:
    p = paragraph._p
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    text_run = OxmlElement("w:r")
    text_run.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text_run, end):
        if node.tag.endswith("fldChar") or node.tag.endswith("instrText"):
            run = OxmlElement("w:r")
            run.append(node)
            p.append(run)
        else:
            p.append(node)


def configure_footer_page_number(section) -> None:
    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clear_runs(paragraph)
    run = paragraph.add_run()
    run.font.name = "宋体"
    run.font.size = Pt(10.5)
    set_rfonts(run, ascii_font="宋体", hansi_font="宋体", east_asia_font="宋体", cs_font="宋体")
    clear_runs(paragraph)
    add_page_field(paragraph)
    for run in paragraph.runs:
        run.font.name = "宋体"
        run.font.size = Pt(10.5)
        set_rfonts(run, ascii_font="宋体", hansi_font="宋体", east_asia_font="宋体", cs_font="宋体")


def add_toc_field(paragraph) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\\\o "1-3" \\\\h \\\\z \\\\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "右键更新域以生成目录"
    separate_run = OxmlElement("w:r")
    separate_run.append(text)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    p = paragraph._p
    for node in (begin, instr, separate, separate_run, end):
        if node.tag.endswith("fldChar") or node.tag.endswith("instrText"):
            run = OxmlElement("w:r")
            run.append(node)
            p.append(run)
        else:
            p.append(node)


def add_hyperlink_run(paragraph, anchor: str, text: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
        rFonts.set(qn(f"w:{key}"), "宋体")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    rPr.append(sz)
    szcs = OxmlElement("w:szCs")
    szcs.set(qn("w:val"), "24")
    rPr.append(szcs)
    t = OxmlElement("w:t")
    t.text = text
    r.append(rPr)
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)


def insert_paragraph_after(paragraph, style_name: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style_name:
        new_para.style = style_name
    return new_para


def cover_fields(source_doc: Document) -> dict[str, str]:
    values = {}
    for row in source_doc.tables[0].rows:
        cells = [normalize(cell.text) for cell in row.cells]
        if len(cells) >= 2 and cells[0]:
            values[cells[0].rstrip("：:")] = cells[1]
    return values


def clean_pdf_text(text: str) -> str:
    lines = []
    for line in text.splitlines():
        s = line.strip()
        if re.fullmatch(r"\d+", s):
            continue
        lines.append(s)
    return "\n".join(lines)


def compact_chinese(text: str) -> str:
    text = normalize(text)
    return re.sub(r"(?<=[\u4e00-\u9fa5])\s+(?=[\u4e00-\u9fa5])", "", text)


def expand_specific_points(item_text: str, task_title: str, target_text: str = "", indicator_text: str = "") -> list[str]:
    topic = re.sub(r"[；;（(].*$", "", item_text).strip("。； ")
    rules: list[tuple[list[str], list[str]]] = [
        (["业务部署架构"], [
            "业务场景与部署层级。",
            "人机角色与协同链路。",
            "架构集成与应用验证。",
        ]),
        (["交互方式", "基础架构", "集成方案"], [
            "交互模式与功能分层。",
            "基础架构与接口集成。",
            "场景部署与联调验证。",
        ]),
        (["模块原型研发"], [
            "原型功能与模块划分。",
            "核心能力研发与组件实现。",
            "原型联调与效果验证。",
        ]),
        (["应用验证"], [
            "验证场景与评估范围。",
            "功能性能联调验证。",
            "应用成效总结与推广分析。",
        ]),
        (["先验知识规则表征", "规则表征"], [
            "知识规则建模与表达。",
            "规则嵌入与模型融合机制。",
            "知识注入效果验证。",
        ]),
        (["知识经验引导", "风险管控决策"], [
            "知识经验约束机制。",
            "风险管控决策策略生成。",
            "决策效果评估与优化。",
        ]),
        (["性能量化评估", "自主趋优"], [
            "性能指标体系与量化评估。",
            "自主趋优机制设计。",
            "量化结果验证与迭代优化。",
        ]),
        (["推理逻辑", "传播机制"], [
            "决策链路与传播表征。",
            "关键因素识别与溯因。",
            "推理机制验证与解释呈现。",
        ]),
        (["全流程可解释", "可解释框架"], [
            "全流程解释框架。",
            "训练应用解释方法。",
            "解释结果验证与呈现。",
        ]),
        (["决策性能优化", "性能优化"], [
            "性能影响因素分析。",
            "优化策略与反馈机制。",
            "性能提升验证。",
        ]),
        (["任务分配", "协作机制"], [
            "态势感知与任务表征。",
            "人机分工与协作策略。",
            "动态交互与机制验证。",
        ]),
        (["主动学习", "主动提问"], [
            "样本选择与知识获取。",
            "主动提问触发机制。",
            "学习反馈与效果验证。",
        ]),
        (["架构", "框架", "体系"], [
            "服务边界与系统分层设计。",
            "数据流与接口规范设计。",
            "部署协同与场景适配验证。",
        ]),
        (["数据", "融合", "接入"], [
            "多源数据对象与接入策略。",
            "异构数据清洗与融合方法。",
            "数据组织与服务机制构建。",
        ]),
        (["并行", "加速", "高通量"], [
            "计算瓶颈与并发能力分析。",
            "并行调度与资源优化方法。",
            "算例验证与加速策略固化。",
        ]),
        (["自动化建模", "建模"], [
            "建模需求与输入输出定义。",
            "特征构造与自动化建模方法。",
            "模型迁移与场景复用验证。",
        ]),
        (["增量", "迭代", "更新"], [
            "样本漂移与性能退化分析。",
            "增量学习与快速迭代策略。",
            "更新评估与回滚机制设计。",
        ]),
        (["可解释", "溯因", "反馈"], [
            "关键影响因素与指标体系。",
            "可解释分析与溯因定位方法。",
            "反馈闭环与模型优化机制。",
        ]),
        (["模块", "研发"], [
            "模块定位与输入输出设计。",
            "关键功能研发与组件封装。",
            "联调测试与支撑效果验证。",
        ]),
    ]
    for keywords, points in rules:
        if any(keyword in topic for keyword in keywords):
            return points
    context = clean_research_text(f"{target_text}；{indicator_text}")
    candidates: list[str] = []
    for clause in split_research_clauses(context):
        clause = clean_research_text(clause)
        if not clause or len(clause) < 6:
            continue
        clause = clause.replace("提出面向", "面向")
        clause = clause.replace("提出", "")
        clause = clause.replace("构建", "")
        clause = clause.replace("建立", "")
        clause = clause.replace("实现", "")
        clause = clause.replace("形成", "")
        clause = clause.replace("增强", "")
        clause = clause.strip("。；，, ")
        if not clause:
            continue
        if topic in clause or clause in topic:
            continue
        if any(word in clause for word in ["论文", "专利", "提交", "发表", "申请"]):
            continue
        candidates.append(clause)
    short_candidates = []
    for clause in candidates:
        parts = re.split(r"[，、]", clause)
        for part in parts:
            part = clean_research_text(part)
            if 4 <= len(part) <= 18 and part not in short_candidates and topic not in part:
                short_candidates.append(part)
    if len(short_candidates) >= 3:
        return [f"{x}。" for x in short_candidates[:3]]
    return [
        "关键问题识别与需求分析。",
        "核心方法设计与机制构建。",
        "验证评估与应用落地。",
    ]


def clean_research_text(text: str) -> str:
    cleaned = compact_chinese(text)
    cleaned = re.sub(r"\s+", "", cleaned)
    cleaned = re.sub(r"[；;，,。\s]*[（(][^）)]*[）)]\s*$", "", cleaned)
    cleaned = cleaned.strip("。；，, ")
    return cleaned


def split_research_clauses(text: str) -> list[str]:
    cleaned = clean_research_text(text)
    if not cleaned:
        return []
    clauses = [seg.strip("。；，, ") for seg in re.split(r"[；。]", cleaned) if seg.strip("。；，, ")]
    if len(clauses) > 1:
        return clauses
    nested = [seg.strip("。；，, ") for seg in re.split(r"，(?=(?:研究|构建|设计|研发|开展|形成|建立|实现|提出|研制|开发|搭建|示范|验证|部署))", cleaned) if seg.strip("。；，, ")]
    if len(nested) > 1 and all(len(seg) >= 8 for seg in nested):
        return nested
    return [cleaned]


def should_merge_leading_clauses(clauses: list[str]) -> bool:
    if len(clauses) < 2:
        return False
    first = clean_research_text(clauses[0])
    second = clean_research_text(clauses[1])
    leading_prefixes = ("基于", "面向", "针对", "围绕", "依托", "结合", "立足", "围绕", "按照")
    action_prefixes = ("设计", "研发", "开展", "构建", "建立", "提出", "形成", "实现", "开发", "研制", "部署", "制定", "搭建", "优化", "探索", "完成")
    if not first.startswith(leading_prefixes):
        return False
    if not second.startswith(action_prefixes):
        return False
    # Treat "基于……框架/方法/机制，设计/研发/开展……" as one complete research item.
    return True


def parse_pdf_source(source_path: Path) -> dict[str, object]:
    reader = PdfReader(str(source_path))
    pages = [page.extract_text() or "" for page in reader.pages]
    full_text = clean_pdf_text("\n".join(pages))
    page1 = clean_pdf_text(pages[0])
    page34 = clean_pdf_text("\n".join(pages[2:4]))

    def search(pattern: str, text: str, flags=re.S, default: str = "") -> str:
        m = re.search(pattern, text, flags)
        return normalize(m.group(1)) if m else default

    project_name = compact_chinese(search(r"项目名称[:：]\s*(.*?)\s*主要承担单位[:：]", page1))
    if not project_name:
        project_name = compact_chinese(search(r"项目名称\s*(.*?)\s*项目\s*总经费", page34))
    main_unit = search(r"主要承担单位[:：]\s*(.*?)\s*主要出资单位[:：]", page1)
    start_time = compact_chinese(search(r"起止时间[:：]\s*(.*?)\s*(?:国\s*家\s*电\s*网\s*公\s*司|$)", page1))
    project_leader = search(r"姓名\s*([^\s]{2,4})\s*单位", page34)

    units = []
    for m in re.finditer(r"\n[1-9]\s+([^\n]+?)\s+\d+\s*$", page34, re.M):
        unit = normalize(m.group(1))
        if unit and unit not in units and "课题" not in unit and "出资单位" not in unit:
            units.append(unit)
    collaborators = [unit for unit in units if unit != main_unit]

    task_leaders = []
    raw_page34 = "\n".join(pages[2:4])
    for m in re.finditer(r"([\u4e00-\u9fa5]{2,4})、\s*\n([\u4e00-\u9fa5]{2,4})\s*\n(?:中国电力|国网信息|国网山东|天津大学)", raw_page34):
        for name in m.groups():
            if name not in task_leaders:
                task_leaders.append(name)
    for m in re.finditer(r"([\u4e00-\u9fa5]{2,4})、([\u4e00-\u9fa5]{1,2})\s*\n([\u4e00-\u9fa5]{1,2})\s*\n(?:中国电力|国网信息|国网山东|天津大学)", raw_page34):
        names = [m.group(1), m.group(2) + m.group(3)]
        for name in names:
            if name not in task_leaders:
                task_leaders.append(name)

    task_pattern = re.compile(
        r"2\.(\d)\s*课题\s*\d+[:：]\s*(.*?)\s*2\.\1\.1\s*研究内容\s*(.*?)\s*2\.\1\.2\s*预期目标\s*(.*?)\s*2\.\1\.3\s*考核指标\s*(.*?)(?=2\.\d\s*课题\s*\d+[:：]|二、分工安排)",
        re.S,
    )

    def split_item_title_and_points(raw: str) -> tuple[str, list[str]]:
        cleaned = clean_research_text(raw)
        clauses = split_research_clauses(cleaned)
        if not clauses:
            return "", []
        title_parts = [clauses[0]]
        start_idx = 1
        if should_merge_leading_clauses(clauses):
            title_parts.append(clauses[1])
            start_idx = 2
        title = clean_research_text("，".join(title_parts))
        points = []
        for point in clauses[start_idx:]:
            point = clean_research_text(point)
            if point and point != title and point not in points:
                points.append(point)
        return title, points

    tasks: list[dict[str, object]] = []
    for match in task_pattern.finditer(full_text):
        no = int(match.group(1))
        title = compact_chinese(match.group(2))
        content_block = match.group(3)
        target = normalize(match.group(4))
        indicator = normalize(match.group(5))

        raw_items = re.findall(r"（\d+）(.*?)(?=（\d+）|$)", content_block, re.S)
        content_items = []
        specific_points_map: dict[str, list[str]] = {}
        details = []
        for raw in raw_items:
            item = clean_research_text(raw)
            units_match = re.search(r"[（(]([^）)]+)[）)]\s*$", compact_chinese(raw))
            unit_text = compact_chinese(units_match.group(1)) if units_match else ""
            title_text, parsed_points = split_item_title_and_points(raw)
            title_text = clean_research_text(title_text)
            if not title_text:
                continue
            content_items.append(title_text)
            points = list(parsed_points)
            if len(points) < 3:
                for extra in expand_specific_points(title_text, title, target, indicator):
                    extra = clean_research_text(extra)
                    if extra and extra not in points and extra != title_text:
                        points.append(extra)
                    if len(points) >= 3:
                        break
            specific_points_map[title_text] = points
            if unit_text:
                details.append(f"{title_text}；承担单位：{unit_text}。")
            else:
                details.append(f"{title_text}。")

        tasks.append(
            {
                "number": no,
                "title": title,
                "overview": (
                    f"任务书将“{title}”设为项目核心课题之一，围绕该课题配置了明确的研究内容、阶段目标与考核要求，"
                    "体现出从基础方法研究到工程应用落地的完整技术路线。"
                ),
                "content_items": content_items,
                "specific_points_map": specific_points_map,
                "target": target,
                "indicator": indicator,
                "details": details,
            }
        )

    return {
        "fields": {
            "项目名称": project_name,
            "主要承担单位": main_unit,
            "起止时间": start_time,
        },
        "participants": {
            "project_leader": project_leader,
            "main_unit": main_unit,
            "collaborators": "、".join(collaborators),
            "collaborators_wrapped": wrap_items_text(collaborators, per_line=2) if collaborators else "",
            "collaborators_body_wrapped": "、".join(collaborators) if collaborators else "",
            "staff": "、".join(task_leaders) if task_leaders else project_leader,
            "staff_wrapped": wrap_staff_text(task_leaders, per_line=4) if task_leaders else project_leader,
            "writers": "、".join(task_leaders[:4]) if task_leaders else project_leader,
        },
        "tasks": tasks,
    }


def load_llm_overrides(path: Path | None) -> dict[str, object]:
    if path is None or not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def merge_task_data(base_tasks: list[dict[str, object]], override_tasks: list[dict[str, object]]) -> list[dict[str, object]]:
    if not override_tasks:
        return base_tasks
    merged: OrderedDict[int, dict[str, object]] = OrderedDict()
    for task in base_tasks:
        merged[int(task.get("number", len(merged) + 1))] = deepcopy(task)
    for idx, task in enumerate(override_tasks, start=1):
        number = int(task.get("number", idx))
        current = merged.get(number, {})
        updated = deepcopy(current)
        updated.update(task)
        if "specific_points_map" in task:
            updated["specific_points_map"] = task["specific_points_map"]
        if "content_items" in task:
            updated["content_items"] = task["content_items"]
        merged[number] = updated
    return [merged[key] for key in sorted(merged)]


def apply_llm_overrides(source_bundle: dict[str, object], overrides: dict[str, object]) -> dict[str, object]:
    if not overrides:
        return source_bundle
    bundle = deepcopy(source_bundle)
    fields = overrides.get("fields")
    if isinstance(fields, dict):
        bundle["fields"].update(fields)
    participants = overrides.get("participants")
    if isinstance(participants, dict):
        bundle["participants"].update(participants)
    override_tasks = overrides.get("tasks")
    if isinstance(override_tasks, list):
        bundle["tasks"] = merge_task_data(bundle.get("tasks", []), override_tasks)
    front_matter = overrides.get("front_matter")
    if isinstance(front_matter, dict):
        bundle["front_matter"] = front_matter
    return bundle


def wrap_items_text(items: list[str], per_line: int = 2, sep: str = "、") -> str:
    lines = []
    for i in range(0, len(items), per_line):
        lines.append(sep.join(items[i:i + per_line]))
    return "\n".join(lines)


def wrap_staff_text(items: list[str], per_line: int = 4, sep: str = "、") -> str:
    lines = []
    for i in range(0, len(items), per_line):
        line = sep.join(items[i:i + per_line])
        if i > 0:
            line = "\t  " + line
        lines.append(line)
    return "\n".join(lines)


def items_by_lines(items: list[str], per_line: int = 2, sep: str = "、") -> list[str]:
    lines = []
    for i in range(0, len(items), per_line):
        lines.append(sep.join(items[i:i + per_line]))
    return lines


def cover_month_text(time_text: str) -> str:
    m = re.search(r"(\d{4})\s*年\s*(\d{1,2})\s*月", time_text)
    if not m:
        return time_text
    return f"{m.group(1)} 年 {int(m.group(2))} 月"


def wrap_cn_title(text: str, first_line: int = 22, next_line: int = 18) -> str:
    text = compact_chinese(text)
    if len(text) <= first_line:
        return text
    # Balance into two lines when the remainder would be too short.
    if len(text) <= first_line + next_line and len(text) - first_line <= 3:
        split = max(1, len(text) // 2)
        return "\n".join([text[:split], text[split:]])
    lines = [text[:first_line]]
    rest = text[first_line:]
    while rest:
        lines.append(rest[:next_line])
        rest = rest[next_line:]
    return "\n".join(lines)


def summarize_task(task: dict[str, object]) -> str:
    return (
        f"综合来看，{task['title']}围绕核心研究目标形成了较完整的技术链条，"
        "既覆盖关键方法和实现路径，也兼顾工程适配与应用验证，"
        "能够为课题后续成果凝练和工程落地提供支撑。"
    )


def expand_to_min_length(text: str, minimum: int = 85, context: str = "") -> str:
    base = compact_chinese(text).strip()
    if len(base) >= minimum:
        return base
    context_label = context or "相关研究内容"
    extras = [
        f"在技术实现过程中，还需要结合{context_label}统筹考虑数据基础、算法设计、人工参与环节、系统部署方式和场景验证要求。",
        "同时，应把任务书中的预期目标和考核指标转化为可执行的技术环节，使研究内容能够对应到模型能力、业务流程和工程应用结果。",
    ]
    enriched = base
    for extra in extras:
        if len(enriched) >= minimum:
            break
        enriched = f"{enriched}{extra}"
    return enriched


def intro_task_heading(title: str) -> str:
    heading = title
    heading = heading.replace("企业级电网智能计算推演", "")
    heading = heading.replace("电网高维异构图模型一体化", "高维异构图")
    heading = heading.replace("基于机器学习自适应建模的", "")
    heading = heading.replace("研究", "")
    heading = heading.replace("技术体系", "体系")
    heading = heading.replace("技术", "")
    heading = heading.replace("及试点应用验证", "与试点验证")
    heading = heading.replace("边云协同智能优化决策功能模块研发与试点应用", "功能模块研发与试点验证")
    heading = heading.strip("：:；; ")
    return heading


def shorten_heading_sentence(text: str) -> str:
    heading = compact_chinese(text).strip("。； ")
    replacements = [
        ("的业务需求与技术边界", "需求与边界"),
        ("业务需求与技术边界", "需求与边界"),
        ("关键方法路线与实现方案", "方法路线与实现方案"),
        ("场景验证与成果沉淀", "场景验证与成果沉淀"),
        ("研究多场景多因素条件下的预训练模型泛化能力提升技术需求与边界", "研究预训练模型泛化提升需求与边界"),
        ("多场景多因素条件下的预训练模型泛化能力提升技术需求与边界", "预训练模型泛化提升需求与边界"),
        ("预训练模型泛化能力提升技术需求与边界", "预训练模型泛化提升需求与边界"),
        ("研究源荷小样本场景下预训练模型参数高效微调技术需求与边界", "研究预训练模型高效微调需求与边界"),
        ("源荷小样本场景下预训练模型参数高效微调技术需求与边界", "预训练模型高效微调需求与边界"),
        ("研究面向源荷小样本场景下预训练模型迁移的场景相似度度量技术需求与边界", "研究场景相似度度量需求与边界"),
        ("面向源荷小样本场景下预训练模型迁移的场景相似度度量技术需求与边界", "场景相似度度量需求与边界"),
        ("研究适用于多时空尺度源荷预测的时域和频域特征构建方法需求与边界", "研究时频特征构建需求与边界"),
        ("研究多时空尺度源荷预测预训练模型输入特征表示学习方法需求与边界", "研究特征表示学习需求与边界"),
        ("研究多时空尺度预训练模型的注意力学习机制构建方法需求与边界", "研究注意力机制构建需求与边界"),
        ("研究多时空尺度预训练模型的主干网络构建方法需求与边界", "研究主干网络构建需求与边界"),
        ("功能需求、仿真需求及业务需求", "功能仿真与业务需求"),
        ("边缘侧与云端", "边云"),
        ("研究柔性互联配电网在不同场景下边云协同互动工作机制", "研究不同场景下的边云协同互动机制"),
        ("协调互动过程中资源主体行为影响关键因素", "协调互动关键因素"),
        ("在不同场景下边缘侧与云端协同互动工作机制", "边云协同互动机制"),
        ("含云服务层、边缘服务层的柔性互联配电网分层多主体边云协同框架", "柔性互联配电网边云协同框架"),
        ("基于深度学习的分布式资源封装体调度特性机理模型", "分布式资源封装体调度机理模型"),
        ("适配多区域子网接线方式与应用场景的自主可控能量装置", "自主可控能量装置"),
        ("基于主导智能体的多区域子网协同控制模型构建方法", "多区域子网协同控制模型方法"),
        ("区域子网内各分布式资源的一致性协同功率分配方法", "区域子网一致性协同功率分配方法"),
        ("基于边云协同的智能优化决策模块研发整体框架", "智能优化决策模块研发框架"),
        ("研究制定统一的数据交互与信息通信接口规范", "研究制定统一数据与通信接口规范"),
        ("研究云边协同联调测试方法", "研究云边联调测试方法"),
        ("研究云端强化学习和边缘侧分布式协同控制算法的验证方法", "研究云边协同算法验证方法"),
        ("研究构建模拟仿真模型，实现配电网潮流计算分析", "研究构建仿真与潮流分析模型"),
        ("针对试点区域配电网柔性互联情况进行资源分析", "开展试点区域配网资源分析"),
        ("梳理总结试点区域配电网拓扑等特征信息", "梳理试点区域配网拓扑特征"),
        ("基于云端强化学习和边缘侧分布式协同算法，实现功能模块的集成开发", "开展云边算法集成功能开发"),
    ]
    for old, new in replacements:
        heading = heading.replace(old, new)
    if heading.startswith("研究") and "需求与边界" in heading:
        heading = heading.replace("研究", "", 1)
    if "，" in heading:
        heading = heading.split("，", 1)[0].strip("。； ")
    heading = heading.replace("、", "")
    heading = heading.replace("，", "")
    return heading.strip("。； ")


def task_status_summary(task: dict[str, object]) -> list[str]:
    items = task.get("content_items", [])
    joined = "、".join(items[:3]) if items else task["title"]
    focus = task_focus_sentence(task)
    indicators = extract_technical_indicators(str(task.get("indicator", "")))
    indicator_sentence = ""
    if indicators:
        indicator_sentence = "从考核指标看，该课题还需重点支撑" + "、".join(indicators[:3]) + "等工程化目标。"
    return [
        expand_to_min_length(
            f"{focus}围绕“{task['title']}”，任务书反映出当前仍存在与该课题相关的关键能力短板，需要针对{joined}等方向开展系统研究。",
            context=task["title"],
        ),
        expand_to_min_length(
            f"从项目整体安排看，该课题既承担基础理论与方法突破任务，也承担后续工程实现和应用验证的前置支撑作用，具有较强的承上启下特征。{indicator_sentence}",
            context=task["title"],
        ),
    ]


def task_content_summary(task: dict[str, object]) -> list[str]:
    items = task.get("content_items", [])
    if items:
        target = compact_target_text(str(task.get("target", ""))).strip("。； ")
        return [
            expand_to_min_length(
                f"该课题的研究内容主要包括：{'；'.join(items[:3])}。这些内容对应任务书提出的“{target}”等目标要求，体现出从关键机制研究到业务能力构建的系统安排。",
                context=task["title"],
            ),
            expand_to_min_length(
                "相关内容共同构成了从关键模型、核心机制到工程实现路径的完整研究链条，是课题成果形成的主体部分，也为后续的模块研发、在线验证和应用推广提供了直接支撑。",
                context=task["title"],
            ),
        ]
    return [
        expand_to_min_length(
            f"该课题围绕“{task['title']}”展开，重点关注关键问题识别、方法构建和落地验证。",
            context=task["title"],
        ),
    ]


def task_route_summary(task: dict[str, object]) -> list[str]:
    items = task.get("content_items", [])
    indicators = extract_technical_indicators(str(task.get("indicator", "")))
    indicator_sentence = ""
    if indicators:
        indicator_sentence = "并通过" + "、".join(indicators[:3]) + "等指标检验技术路线的有效性"
    if items:
        first = items[0]
        last = items[-1]
        return [
            expand_to_min_length(
                f"从技术路线看，该课题按照“问题分析 - 关键方法设计 - 系统实现与验证”的路径推进，并以“{first}”作为切入点，以“{last}”作为落地验证的重要抓手，{indicator_sentence}。",
                context=task["title"],
            ),
            expand_to_min_length(
                "这一技术路线兼顾理论创新、系统实现和场景适配，有利于形成可复用、可验证、可推广的阶段性成果，并使研究结果能够真正服务于电网运行分析和决策业务。",
                context=task["title"],
            ),
        ]
    return [
        expand_to_min_length(
            "该课题总体遵循问题识别、方法研究、系统实现和应用验证的技术路线，逐步完成关键能力构建。",
            context=task["title"],
        ),
    ]


def compact_target_text(text: str) -> str:
    normalized = compact_chinese(text).strip("。； ")
    normalized = re.sub(r"\s+", "", normalized)
    return normalized


def extract_technical_indicators(indicator_text: str) -> list[str]:
    text = compact_target_text(indicator_text)
    patterns = [
        r"精度相较常规机器学习算法提升大于\d+%",
        r"识别准确率大于\d+%",
        r"模型训练收敛时间减少\d+%以上",
        r"人机决策冲突率.*?降低\d+%以上",
        r"支撑电网计算节点规模不低于\d+个",
        r"薄弱环节分析定位时间小于\d+s",
        r"负荷转供策略生成时间小于\d+s",
        r"故障恢复处置平均时间降低\d+%",
        r"考虑预想事故数目不低于\d+个",
        r"提供停电概率、停电时间、停电电量、停电户数等\d+类指标",
    ]
    hits: list[str] = []
    for pattern in patterns:
        for m in re.finditer(pattern, text):
            value = m.group(0)
            if value not in hits:
                hits.append(value)
    return hits[:4]


def project_business_context(project_name: str, tasks: list[dict[str, object]]) -> str:
    corpus = compact_target_text(project_name + "；" + "；".join(
        [str(task.get("title", "")) + "；" + str(task.get("target", "")) for task in tasks]
    ))
    parts: list[str] = []
    if "可靠性" in corpus:
        parts.append("地区电网供电可靠性评估、停电风险量化和薄弱环节辨识")
    if "热稳" in corpus or "热稳定" in corpus:
        parts.append("源荷不确定条件下的热稳定态势感知与关键断面预测")
    if "风险预警" in corpus or "风险管控" in corpus:
        parts.append("电网运行风险预警、安全校正控制和风险管控策略生成")
    if "转供" in corpus or "应急处置" in corpus:
        parts.append("应急处置预案在线生成与负荷转供策略优选")
    if "可解释" in corpus or "溯因" in corpus:
        parts.append("模型决策过程溯因分析与人机互信增强")
    if "主动学习" in corpus or "主动提问" in corpus:
        parts.append("小样本条件下的主动学习与知识获取")
    if not parts:
        parts.append("电网业务场景下的模型构建、决策分析与工程验证")
    return "、".join(parts)


def task_focus_sentence(task: dict[str, object]) -> str:
    title = compact_target_text(str(task.get("title", "")))
    if "可靠性分析" in title and "协作机制" in title:
        return "该课题聚焦地区电网可靠性分析场景下的人机混合增强智能架构、任务分配和交互协作机制，重点解决人工智能置信度不足时人工及时介入、协同决策链条闭环以及小样本条件下模型持续学习等问题。"
    if "可解释" in title or "溯因" in title:
        return "该课题聚焦地区电网可靠性分析结果的可解释建模与溯因分析，重点解决深度学习模型决策链条不透明、运行人员难以理解模型依据以及解释结果难以反馈优化模型的问题。"
    if "知识引导学习" in title or "风险管控决策" in title:
        return "该课题聚焦知识引导下的混合增强风险管控决策技术，重点解决电网人工智能模型训练样本依赖高、决策逻辑与专家经验冲突以及模型长期运行持续趋优能力不足等问题。"
    if "模块研发" in title or "应用验证" in title:
        return "该课题聚焦人机协同混合增强智能决策模块研发及应用验证，重点解决可靠性快速评估、薄弱点辨识、风险预警、热稳态势感知和转供策略生成等场景中的工程集成与在线应用问题。"
    return f"该课题围绕“{task['title']}”展开，重点解决任务书中提出的关键技术问题，并形成从方法研究到工程验证的完整实现路径。"


def item_focus_sentence(task: dict[str, object], item: str) -> str:
    topic = compact_target_text(item)
    if "业务部署架构" in topic:
        return "本节围绕地区电网可靠性分析的人机混合增强智能部署架构展开，重点说明业务分析链路、计算部署层级、人工研判环节和智能算法模块之间的协同组织方式。"
    if "任务分配" in topic or "协作机制" in topic:
        return "本节围绕实时态势评价驱动的人机协作机制展开，重点说明在可靠性评估、薄弱点辨识和风险处置过程中如何根据模型置信度、任务紧急程度和业务复杂度动态分配人机职责。"
    if "主动学习" in topic or "主动提问" in topic:
        return "本节围绕机器向人的主动学习与主动提问方法展开，重点说明在标注样本有限、工况复杂多变的地区电网场景下，如何通过人机交互持续提升模型学习效率与分析精度。"
    if "推理逻辑" in topic or "传播机制" in topic:
        return "本节围绕模型决策过程的推理逻辑与信息传播机制展开，重点说明如何从特征传递、状态演化和决策触发链条中还原地区电网可靠性分析结果的形成依据。"
    if "全流程可解释" in topic or "可解释框架" in topic:
        return "本节围绕模型构建与应用的全流程可解释框架展开，重点说明训练、推理、结果展示和人工反馈各阶段的解释信息如何贯通。"
    if "决策性能优化" in topic:
        return "本节围绕基于模型可解释机制的决策性能优化方法展开，重点说明如何将解释结果反馈到特征选择、参数调整和决策校正过程中。"
    if "规则表征" in topic or "嵌入方法" in topic:
        return "本节围绕先验知识规则的表征与嵌入方法展开，重点说明电网运行规则、专家经验和安全约束如何转化为模型可识别、可调用的知识表达。"
    if "风险管控决策" in topic:
        return "本节围绕知识经验引导下的风险管控决策技术展开，重点说明如何将专家知识、运行方式和风险边界约束到模型决策过程中，提升策略生成的安全性和可接受性。"
    if "性能量化评估" in topic or "自主趋优" in topic:
        return "本节围绕混合增强智能性能量化评估与自主趋优技术展开，重点说明如何对模型效果、人机协作质量和策略执行成效进行量化评价，并形成持续优化闭环。"
    if "模块原型研发" in topic:
        return "本节围绕人机协同混合增强智能决策模块原型研发展开，重点说明模块功能组成、算法能力封装、业务接口适配和原型系统联调实现路径。"
    if "交互方式" in topic or "集成方案" in topic:
        return "本节围绕决策模块的交互方式、基础架构与集成方案展开，重点说明分析界面、算法服务、知识组件和电网业务系统之间的集成关系。"
    if "应用验证" in topic:
        return "本节围绕混合增强智能决策模块在地区电网中的应用验证展开，重点说明典型业务场景选择、验证指标设计、在线联调方式和应用成效评估思路。"
    return f"本节围绕“{item}”展开，重点说明其在电网业务场景中的实现思路、关键技术环节和验证方式。"


def point_implementation_paragraphs(task: dict[str, object], item: str, point: str) -> list[str]:
    point_text = compact_target_text(point)
    item_text = compact_target_text(item)
    indicators = extract_technical_indicators(str(task.get("indicator", "")))
    indicator_sentence = ""
    if indicators:
        indicator_sentence = "结合任务书考核指标，相关实现还需要重点支撑" + "、".join(indicators[:3]) + "等要求。"
    base_context = item
    if "部署层级" in point_text:
        paragraphs = [
            "在业务场景与部署层级设计方面，应结合地区电网供电可靠性快速评估、停电风险量化、薄弱环节辨识等典型业务链路，明确数据接入层、状态感知层、模型分析层、人工研判层和决策服务层的职责边界，形成面向分析与决策一体化的部署结构。",
            f"具体实现上，需要将可靠性指标计算、故障场景生成、风险态势评估与人工复核节点进行分层部署，使人工专家能够在模型置信度不足、事故组合复杂或策略后果敏感的环节及时介入。{indicator_sentence}",
        ]
    elif "协同链路" in point_text or "分工与协作" in point_text:
        paragraphs = [
            "在人机角色与协同链路设计方面，需要明确模型负责的快速筛查、候选方案生成和风险排序能力，以及人工负责的异常识别、知识校核、策略裁决和结果确认职责，形成可追踪、可回溯的人机协作闭环。",
            f"具体到地区电网场景，应围绕可靠性评估、风险预警和安全管控策略生成等业务，建立基于置信度、解释结果和运行边界的任务切换机制，使人机协同既保持分析效率，又保证关键决策安全可控。{indicator_sentence}",
        ]
    elif "主动学习" in item_text or "知识获取" in point_text or "主动提问" in point_text:
        paragraphs = [
            "在主动学习与主动提问实现方面，需要围绕样本稀缺而工况多变的地区电网场景，优先挑选对模型边界最敏感、对风险识别最关键的样本提交人工标注，并通过问题模板、知识问答或交互确认机制补足模型难以直接学习的专业知识。",
            f"通过将人工反馈转化为增量训练样本、规则约束或特征修正信息，可逐步提升模型在小样本条件下的泛化能力和业务适应性，支撑任务书中对主动学习精度提升和持续学习能力的要求。{indicator_sentence}",
        ]
    elif "解释框架" in point_text or "解释方法" in point_text or "溯因" in point_text:
        paragraphs = [
            "在可解释与溯因实现方面，应从输入特征贡献、网络中间表征、输出结果敏感性和规则触发依据等多个层面构建解释链条，使运行人员能够理解模型为何给出某一可靠性判断或风险预警结果。",
            f"具体实现上，需要将训练阶段的样本贡献分析、推理阶段的重要特征定位以及应用阶段的结果回溯展示统一起来，并把解释信息反向用于参数修正、特征优化和策略校正，形成面向人机互信的闭环机制。{indicator_sentence}",
        ]
    elif "规则建模" in point_text or "规则嵌入" in point_text or "知识注入" in point_text:
        paragraphs = [
            "在知识规则表征与嵌入方面，需要把电网运行规程、安全约束、调度经验和典型事故处置原则抽象为可计算的规则表达，并设计与机器学习模型兼容的嵌入方式，使模型在训练和推理阶段都能利用显式知识。",
            f"通过将知识规则与数据驱动模型联合建模，可减少对大规模标注样本的依赖，提升风险管控决策的安全边界一致性，并降低人机决策逻辑冲突。{indicator_sentence}",
        ]
    elif "性能量化" in point_text or "趋优" in point_text or "性能提升" in point_text:
        paragraphs = [
            "在性能量化评估与持续优化方面，需要面向可靠性评估准确性、风险识别效率、人机协作质量和策略执行效果构建多维指标体系，既评价模型本身性能，也评价人工介入后的协同增益。",
            f"基于评估结果，应建立自动调参与知识更新协同机制，对误判场景、高风险样本和策略偏差进行闭环修正，推动模型能力持续趋优。{indicator_sentence}",
        ]
    elif "模块" in item_text or "原型" in point_text or "集成" in point_text or "验证" in point_text:
        paragraphs = [
            "在模块研发与工程集成方面，需要围绕地区电网可靠性评估、热稳态势感知、风险预警、应急处置和转供策略生成等典型场景，完成算法能力、知识组件、交互界面和业务接口的一体化封装。",
            f"具体实现上，应打通数据接入、模型服务、解释展示、人工交互和策略输出链路，通过联调测试验证原型模块对业务系统的支撑能力，并围绕节点规模、在线分析时效和策略生成效率等指标开展应用评估。{indicator_sentence}",
        ]
    else:
        paragraphs = [
            f"围绕“{point.strip('。； ')}”这一实现要点，需要结合{task.get('title', '')}的业务目标，细化算法设计、系统实现和工程验证之间的衔接关系，使研究内容能够直接支撑地区电网相关场景的分析与决策需求。",
            f"在实施过程中，应同步考虑数据质量、业务规则、模型解释性和应用验证要求，将研究方法落到可部署、可评估、可优化的技术环节中。{indicator_sentence}",
        ]
    return [expand_to_min_length(p, minimum=95, context=base_context) for p in paragraphs]


def item_section_paragraphs(task: dict[str, object], item: str) -> list[str]:
    indicators = extract_technical_indicators(str(task.get("indicator", "")))
    indicator_sentence = ""
    if indicators:
        indicator_sentence = "结合任务书考核指标，本节研究需要同时关注" + "、".join(indicators[:3]) + "等可验证结果。"
    paragraphs = [
        item_focus_sentence(task, item),
        f"从实现路径看，该研究内容需要把任务书中的预期目标转化为可执行的技术环节，围绕数据组织、模型构建、知识注入、人机交互、系统部署和场景验证等方面逐步展开。{indicator_sentence}",
    ]
    return [expand_to_min_length(p, minimum=95, context=item) for p in paragraphs]


def task_overview_paragraphs(task: dict[str, object]) -> list[str]:
    indicators = extract_technical_indicators(str(task.get("indicator", "")))
    indicator_sentence = ""
    if indicators:
        indicator_sentence = "结合考核指标，课题研究还需要重点满足" + "、".join(indicators[:3]) + "等验证要求。"
    paragraphs = [
        task_focus_sentence(task),
        f"结合任务书预期目标，本课题需要把研究内容落实到可执行的技术方案中，围绕模型方法、知识机制、人机交互、系统实现和业务验证形成闭环。{indicator_sentence}",
    ]
    return [expand_to_min_length(p, minimum=100, context=str(task.get('title', ''))) for p in paragraphs]


def parse_participants(source_doc: Document) -> dict[str, str]:
    table = source_doc.tables[1]
    project_leader = normalize(table.rows[3].cells[2].text)
    work_units = []
    for row in table.rows[8:14]:
        unit = normalize(row.cells[2].text)
        if unit and unit not in work_units:
            work_units.append(unit)
    collaborators = [unit for unit in work_units[1:] if unit]

    task_leaders = []
    for row in table.rows[17:22]:
        names = normalize(row.cells[14].text)
        for name in re.split(r"[、，, ]+", names):
            if name and name not in task_leaders:
                task_leaders.append(name)

    return {
        "project_leader": project_leader,
        "main_unit": work_units[0] if work_units else "",
        "collaborators_list": collaborators,
        "collaborators": "、".join(collaborators),
        "collaborators_wrapped": wrap_items_text(collaborators, per_line=2),
        "collaborators_body_wrapped": "、".join(collaborators),
        "staff": "、".join(task_leaders),
        "staff_wrapped": wrap_items_text(task_leaders, per_line=4),
        "writers": "、".join(task_leaders[:4]) if task_leaders else project_leader,
    }


def english_title(chinese_title: str) -> str:
    normalized = compact_chinese(chinese_title)
    title_map = {
        "面向企业级电网智能计算推演的共性服务关键技术研究": "Research on Key Common Service Technologies for Enterprise-Level Intelligent Power Grid Computation and Simulation",
        "基于知识增强的电网运行可信智能决策基础理论研究": "Basic Theoretical Research on Knowledge-Enhanced Trusted Intelligent Decision-Making for Power Grid Operation",
        "人机协同的混合增强智能技术及其在地区电网可靠性分析中的应用": "Human-Machine Collaborative Hybrid-Augmented Intelligence Technology and Its Application in Regional Power Grid Reliability Analysis",
    }
    if normalized in title_map:
        return title_map[normalized]

    phrase_map = [
        ("企业级电网智能计算推演", "Enterprise-Level Intelligent Power Grid Computation and Simulation"),
        ("电网运行可信智能决策", "Trusted Intelligent Decision-Making for Power Grid Operation"),
        ("人机协同", "Human-Machine Collaboration"),
        ("混合增强智能", "Hybrid-Augmented Intelligence"),
        ("地区电网可靠性分析", "Regional Power Grid Reliability Analysis"),
        ("可靠性分析", "Reliability Analysis"),
        ("中的应用", "Application"),
        ("及其在", "and Its Application in"),
        ("机器学习自适应建模", "Machine Learning-Based Adaptive Modeling"),
        ("多时空推演", "Multi-Temporal and Spatial Simulation"),
        ("知识增强智能可信决策理论", "Knowledge-Enhanced Trusted Intelligent Decision Theory"),
        ("知识增强智能可信决策", "Knowledge-Enhanced Trusted Intelligent Decision-Making"),
        ("智能可信决策理论", "Trusted Intelligent Decision Theory"),
        ("共性服务关键技术", "Key Common Service Technologies"),
        ("共性服务", "Common Services"),
        ("知识增强", "Knowledge Enhancement"),
        ("智能计算推演", "Intelligent Computation and Simulation"),
        ("电网运行", "Power Grid Operation"),
        ("基础理论", "Basic Theory"),
        ("理论研究", "Theoretical Research"),
        ("关键技术", "Key Technologies"),
        ("智能决策", "Intelligent Decision-Making"),
        ("电网", "Power Grid"),
    ]

    def translate_segment(text: str) -> str:
        result = compact_chinese(text)
        for zh, en in phrase_map:
            result = result.replace(zh, en)
        result = result.replace("、", " and ")
        result = result.replace("与", " and ")
        result = result.replace("及", " and ")
        result = re.sub(r"\s+", " ", result).strip(" ,.-")
        return result

    m = re.match(r"^面向(.+)的(.+?)(关键技术研究|技术研究|基础理论研究|研究)$", normalized)
    if m:
        context = translate_segment(m.group(1))
        subject = translate_segment(m.group(2))
        suffix = m.group(3)
        if suffix == "关键技术研究":
            return f"Research on {subject} for {context}"
        if suffix == "基础理论研究":
            return f"Basic Theoretical Research on {subject} for {context}"
        return f"Research on {subject} for {context}"

    m = re.match(r"^基于(.+)的(.+?)(基础理论研究|关键技术研究|技术研究|研究)$", normalized)
    if m:
        base = translate_segment(m.group(1))
        subject = translate_segment(m.group(2))
        suffix = m.group(3)
        if base == "Knowledge Enhancement":
            base = "Knowledge-Enhanced"
            connector = ""
        else:
            connector = f" Based on {base}"
        if suffix == "基础理论研究":
            return f"Basic Theoretical Research on {subject}{connector}"
        if suffix == "关键技术研究":
            return f"Research on Key Technologies for {subject}{connector}"
        return f"Research on {subject}{connector}"

    m = re.match(r"^(.+?)的(.+?)及其在(.+?)中的应用$", normalized)
    if m:
        actor = translate_segment(m.group(1))
        subject = translate_segment(m.group(2))
        application = translate_segment(m.group(3))
        if actor == "Human-Machine Collaboration":
            actor = "Human-Machine Collaborative"
        return f"{actor} {subject} and Its Application in {application}"

    generic = translate_segment(normalized)
    if generic and re.search(r"[A-Za-z]", generic):
        return generic
    return "Technical Report"


def replace_cover(output: Document, source_doc: Document) -> None:
    fields = cover_fields(source_doc)
    participants = parse_participants(source_doc)
    time_text = fields.get("起止时间", "")
    time_compact = (
        time_text.replace("年", ".").replace("月", "").replace("至", "--").replace(" ", "")
        if time_text
        else ""
    )
    replacements = {
        "技术报告中文名称": fields.get("项目名称", ""),
        "Name of Technical Report": fields.get("英文标题") or english_title(fields.get("项目名称", "")),
        "《科研项目名称》总报告/系列报告（阿拉伯数字）": f"《{fields.get('项目名称', '技术报告')}》技术总报告",
        "中国电力科学研究院有限公司": participants.get("main_unit", fields.get("主要承担单位", "")),
        "合作（协作）单位": participants.get("collaborators", ""),
        "20××年××月": time_text,
    }
    for paragraph in output.paragraphs:
        text = normalize(paragraph.text)
        if not text:
            continue
        if text == participants.get("collaborators", ""):
            text = participants.get('collaborators_wrapped', participants.get('collaborators', ''))
        elif text.startswith("工作单位："):
            text = f"工作单位： {participants.get('main_unit', '')}"
        elif text.startswith("合作（协作）单位"):
            text = f"合作（协作）单位： {participants.get('collaborators_body_wrapped', participants.get('collaborators', ''))}"
        elif text.startswith("工作时间："):
            text = f"工作时间：{time_compact}"
        elif text.startswith("项目负责人："):
            text = f"项目负责人： {participants.get('project_leader', '')}"
        elif text.startswith("工作人员："):
            text = f"工作人员：{participants.get('staff_wrapped', participants.get('staff', ''))}"
        elif text.startswith("报告编写："):
            text = f"报告编写：{participants.get('writers', '')}"
        elif normalize(text) == "中国电力科学研究院有限公司":
            text = participants.get("main_unit", "")
        elif text.startswith("报告审核") or text.startswith("学术技术分委会：") or text.startswith("部门领导：") or text.startswith("科技创新部/咨询业务部（格式审核）：") or text.startswith("院学术技术委员会：") or text.startswith("报告批准") or text.startswith("副院长/总工程师/院长助理/副总工程师："):
            text = re.sub(r"[:：].*$", "：", text)
        elif text in replacements:
            text = replacements[text]
        else:
            continue
        sample = first_nonempty_run(paragraph)
        clear_runs(paragraph)
        run = paragraph.add_run(text)
        if sample is not None:
            copy_run_font(run, sample)


def replace_cover_from_bundle(output: Document, source_bundle: dict[str, object]) -> None:
    fields = source_bundle["fields"]
    participants = source_bundle["participants"]
    time_text = fields.get("起止时间", "")
    cover_date = cover_month_text(time_text)
    time_compact = (
        time_text.replace("年", ".").replace("月", "").replace("至", "--").replace(" ", "")
        if time_text
        else ""
    )
    collaborator_lines = items_by_lines(participants.get("collaborators_list", []), per_line=2)
    replacements = {
        "技术报告中文名称": wrap_cn_title(fields.get("项目名称", ""), 30, 24),
        "Name of Technical Report": fields.get("英文标题") or english_title(fields.get("项目名称", "")),
        "《科研项目名称》总报告/系列报告（阿拉伯数字）": f"《{fields.get('项目名称', '技术报告')}》技术总报告",
        "中国电力科学研究院有限公司": participants.get("main_unit", fields.get("主要承担单位", "")),
        "合作（协作）单位": participants.get("collaborators", ""),
        "20××年××月": cover_date,
    }
    for idx, paragraph in enumerate(output.paragraphs):
        text = normalize(paragraph.text)
        if not text:
            continue
        if idx == 9:
            text = wrap_cn_title(fields.get("项目名称", ""), 30, 24)
        elif idx == 10:
            text = english_title(fields.get("项目名称", ""))
        elif idx == 11:
            text = f"《{fields.get('项目名称', '技术报告')}》技术总报告"
        elif idx == 14:
            text = participants.get("main_unit", fields.get("主要承担单位", ""))
        elif idx == 15 and re.search(r"\d{4}\s*年", text):
            text = cover_date
        elif idx == 16 and text.startswith("工作单位"):
            text = f"工作单位：{participants.get('main_unit', '')}"
        elif idx == 17 and not text.startswith("工作时间") and not text.startswith("项目负责人"):
            text = collaborator_lines[0] if len(collaborator_lines) > 0 else ""
        elif idx == 18 and not text.startswith("工作时间") and not text.startswith("项目负责人"):
            text = collaborator_lines[1] if len(collaborator_lines) > 1 else ""
        elif idx == 19 and not text.startswith("工作时间") and not text.startswith("项目负责人"):
            text = collaborator_lines[2] if len(collaborator_lines) > 2 else ""
        elif idx == 21 and text.startswith("工作时间"):
            text = f"工作时间：{time_compact}"
        elif idx == 22 and text.startswith("项目负责人："):
            text = f"项目负责人： {participants.get('project_leader', '')}"
        elif idx == 23 and text.startswith("工作人员："):
            text = f"工作人员：{participants.get('staff_wrapped', participants.get('staff', ''))}"
        elif idx == 24 and text.startswith("报告编写："):
            text = f"报告编写：{participants.get('writers', '')}"
        elif idx == 17:
            text = participants.get("main_unit", fields.get("主要承担单位", ""))
        elif idx == 18:
            text = ""
        elif idx == 19:
            text = cover_date
        elif idx == 21:
            text = f"工作单位：{participants.get('main_unit', '')}"
        elif idx == 22:
            text = collaborator_lines[0] if len(collaborator_lines) > 0 else participants.get("main_unit", "")
        elif idx == 23:
            text = collaborator_lines[1] if len(collaborator_lines) > 1 else ""
        elif idx == 24:
            text = collaborator_lines[2] if len(collaborator_lines) > 2 else ""
        elif idx in {25, 26}:
            text = ""
        elif text == participants.get("collaborators", ""):
            text = participants.get('collaborators_wrapped', participants.get('collaborators', ''))
        elif text.startswith("工作单位："):
            text = f"工作单位：{participants.get('main_unit', '')}"
        elif text.startswith("合作（协作）单位"):
            text = ""
        elif text.startswith("工作时间："):
            text = f"工作时间：{time_compact}"
        elif text.startswith("项目负责人："):
            text = f"项目负责人： {participants.get('project_leader', '')}"
        elif text.startswith("工作人员："):
            text = f"工作人员：{participants.get('staff_wrapped', participants.get('staff', ''))}"
        elif text.startswith("报告编写："):
            text = f"报告编写：{participants.get('writers', '')}"
        elif normalize(text) == "中国电力科学研究院有限公司":
            text = participants.get("main_unit", "")
        elif text.startswith("报告审核") or text.startswith("学术技术分委会：") or text.startswith("部门领导：") or text.startswith("科技创新部/咨询业务部（格式审核）：") or text.startswith("院学术技术委员会：") or text.startswith("报告批准") or text.startswith("副院长/总工程师/院长助理/副总工程师："):
            text = re.sub(r"[:：].*$", "：", text)
        elif text in replacements:
            text = replacements[text]
        else:
            continue
        sample = first_nonempty_run(paragraph)
        clear_runs(paragraph)
        run = paragraph.add_run(text)
        if sample is not None:
            copy_run_font(run, sample)


def replace_fixed_front_matter_fields(output: Document, source_bundle: dict[str, object]) -> None:
    fields = source_bundle["fields"]
    participants = source_bundle["participants"]
    time_text = fields.get("起止时间", "")
    cover_date = cover_month_text(time_text)
    time_compact = (
        time_text.replace("年", ".").replace("月", "").replace("至", "--").replace(" ", "")
        if time_text
        else ""
    )

    for idx, paragraph in enumerate(output.paragraphs):
        text = normalize(paragraph.text)
        if not text:
            continue

        new_text = None
        if idx == 9:
            new_text = compact_chinese(fields.get("项目名称", ""))
        elif idx == 10:
            new_text = fields.get("英文标题") or english_title(fields.get("项目名称", ""))
        elif idx == 11:
            new_text = f"《{fields.get('项目名称', '技术报告')}》技术总报告"
        elif idx == 14:
            new_text = participants.get("main_unit", fields.get("主要承担单位", ""))
        elif idx == 15:
            new_text = cover_date
        elif text.startswith("工作单位："):
            new_text = f"工作单位：{participants.get('main_unit', '')}"
        elif text.startswith("工作时间："):
            new_text = f"工作时间：{time_compact}"
        elif text.startswith("项目负责人："):
            new_text = f"项目负责人： {participants.get('project_leader', '')}"
        elif text.startswith("工作人员："):
            new_text = f"工作人员：{participants.get('staff_wrapped', participants.get('staff', ''))}"
        elif text.startswith("报告编写："):
            new_text = f"报告编写：{participants.get('writers', '')}"
        elif idx == 27 and normalize(text):
            new_text = participants.get("main_unit", fields.get("主要承担单位", ""))

        if new_text is None:
            continue

        sample = first_nonempty_run(paragraph)
        clear_runs(paragraph)
        run = paragraph.add_run(new_text)
        if sample is not None:
            copy_run_font(run, sample)
        if idx in {9, 10, 11, 14, 15}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def postprocess_cover_layout(doc: Document) -> None:
    centered = {9, 10, 11, 17, 18, 19}
    signature_titles = {31, 37}
    for idx, paragraph in enumerate(doc.paragraphs[:40]):
        if idx in centered:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = None
        if idx == 9:
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(12)
        if idx == 10:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            for run in paragraph.runs:
                run.font.size = Pt(14)
                set_rfonts(run, ascii_font="Times New Roman", hansi_font="Times New Roman", east_asia_font="Times New Roman", cs_font="Times New Roman")
        if idx == 11:
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(6)
        if idx == 17:
            paragraph.paragraph_format.space_before = Pt(48)
            paragraph.paragraph_format.space_after = Pt(6)
        if idx == 19:
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(0)
        if idx in {3, 4, 5}:
            ppr = paragraph._p.get_or_add_pPr()
            ind = ppr.find(qn("w:ind"))
            if ind is not None:
                ppr.remove(ind)
            tabs = ppr.find(qn("w:tabs"))
            if tabs is not None:
                ppr.remove(tabs)
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.left_indent = None
            paragraph.paragraph_format.right_indent = Pt(54)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if idx in signature_titles:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.bold = True
        if idx in {21, 22, 24, 25, 26, 27, 30, 31, 32, 33, 34, 35, 37, 38}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.left_indent = None
            paragraph.paragraph_format.right_indent = None
            paragraph.paragraph_format.line_spacing = 1.0
        if idx in {21, 24, 25, 26, 27, 30, 31, 37}:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(3)
        if idx == 22:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
        if idx == 36:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 0.8
        if idx in {32, 33, 34, 35, 38}:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(3)
        if idx == 12:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.left_indent = None
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.line_spacing = 1.5
        if idx == 14:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = None
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.line_spacing = 1.5
        if idx in {15, 16, 17, 18}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Pt(79.8)
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.5
        if idx == 25:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = None
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.line_spacing = 1.5
        if idx in {26, 27, 28, 29, 30}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = None
            paragraph.paragraph_format.first_line_indent = Pt(70.05)
            paragraph.paragraph_format.line_spacing = 1.5
        if idx in {32, 33}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Pt(29.75)
            paragraph.paragraph_format.first_line_indent = Pt(42.05)
            paragraph.paragraph_format.line_spacing = 1.5
    # Remove blank template spacer paragraphs so the whole sign-off block stays on one page.
    for idx in [36, 23, 20, 18, 16, 15, 14, 13, 12]:
        if idx < len(doc.paragraphs) and not normalize(doc.paragraphs[idx].text):
            element = doc.paragraphs[idx]._element
            element.getparent().remove(element)
    for paragraph in doc.paragraphs[:40]:
        if normalize(paragraph.text).startswith("工作单位："):
            paragraph.paragraph_format.page_break_before = True
            break
    for paragraph in doc.paragraphs[:80]:
        text = normalize(paragraph.text)
        if text in {"内容摘要", "ABSTRACT", "目 录"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.left_indent = None
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.5
        elif text.startswith("工作单位："):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = None
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.line_spacing = 1.5


def clear_body_after(doc: Document, target_text: str) -> None:
    body = doc.element.body
    children = list(body.iterchildren())
    start = None
    for idx, child in enumerate(children):
        texts = []
        for node in child.iter():
            if node.tag.endswith("}t") and node.text:
                texts.append(node.text)
        if normalize("".join(texts)) == target_text:
            start = idx
            break
    if start is None:
        raise ValueError(target_text)
    sect_pr = children[-1] if children and children[-1].tag.endswith("}sectPr") else None
    for child in children[start:]:
        if sect_pr is not None and child is sect_pr:
            break
        body.remove(child)


def find_template_paragraph(template: Document, *, exact_text: str | None = None, startswith: str | None = None, style_name: str | None = None):
    for paragraph in template.paragraphs:
        text = normalize(paragraph.text)
        if not text:
            continue
        if exact_text is not None and text != exact_text:
            continue
        if startswith is not None and not text.startswith(startswith):
            continue
        if style_name is not None and paragraph.style.name != style_name:
            continue
        return paragraph
    raise ValueError(f"Template paragraph not found: exact_text={exact_text!r}, startswith={startswith!r}, style_name={style_name!r}")


def replace_front_matter_from_reference(doc: Document, ref_doc: Document, start_marker: str = "内容摘要") -> None:
    def find_marker_child(document: Document, marker: str):
        for child in document.element.body.iterchildren():
            texts = []
            for node in child.iter():
                if node.tag.endswith("}t") and node.text:
                    texts.append(node.text)
            if normalize("".join(texts)) == marker:
                return child
        raise ValueError(marker)

    body = doc.element.body
    anchor = find_marker_child(doc, start_marker)
    children = list(body.iterchildren())
    sect_pr = children[-1] if children and children[-1].tag.endswith("}sectPr") else None
    for child in children:
        if child is anchor or (sect_pr is not None and child is sect_pr):
            break
        body.remove(child)

    ref_body = ref_doc.element.body
    ref_anchor = find_marker_child(ref_doc, start_marker)
    ref_children = list(ref_body.iterchildren())
    front_children = []
    for child in ref_children:
        if child is ref_anchor or child.tag.endswith("}sectPr"):
            break
        front_children.append(deepcopy(child))

    for child in front_children:
        anchor.addprevious(child)


def extract_tasks(source_doc: Document) -> list[dict[str, object]]:
    text = normalize(source_doc.tables[2].cell(0, 0).text)
    pattern = re.compile(
        r"2\.(\d)课题\d+：([^。]+?)2\.\1\.1研究内容(.*?)2\.\1\.2预期目标(.*?)2\.\1\.3考核指标(.*?)(?=2\.\d课题\d+：|$)"
    )
    task_map: OrderedDict[str, dict[str, object]] = OrderedDict()
    for match in pattern.finditer(text):
        no = int(match.group(1))
        content_text = normalize(match.group(3))
        content_items = [normalize(item) for item in re.findall(r"（\d+）([^；。]+(?:；（[^）]+）)?)", content_text)]
        if not content_items:
            content_items = [normalize(item) for item in re.findall(r"([^；。]+?；（[^）]+）)", content_text)]
        if not content_items:
            content_items = [
                normalize(part)
                for part in re.split(r"(?<=）)\s+", content_text)
                if normalize(part)
            ]
        if not content_items or (len(content_items) == 1 and content_items[0] == content_text):
            content_items = [normalize(part) for part in re.split(r"[；。]", content_text) if normalize(part)]
        task_map[str(no)] = {
            "number": no,
            "title": match.group(2).strip(" ：:"),
            "overview": "",
            "content": content_text,
            "content_items": content_items,
            "target": normalize(match.group(4)),
            "indicator": normalize(match.group(5)),
            "details": [],
        }

    for row in source_doc.tables[3].rows[1:]:
        cells = [normalize(c.text) for c in row.cells]
        if len(cells) < 4 or not cells[0]:
            continue
        m = re.match(r"课题(\d+)：(.+)", cells[0])
        if not m:
            continue
        no = m.group(1)
        task = task_map.setdefault(
            no,
            {
                "number": int(no),
                "title": m.group(2),
                "overview": "",
                "content": "",
                "content_items": [],
                "target": "",
                "indicator": "",
                "details": [],
            },
        )
        task["title"] = m.group(2)
        detail = f"{cells[1]}；承担单位：{cells[2]}；经费：{cells[3]}万元。"
        task["details"].append(detail)

    for task in task_map.values():
        task["overview"] = (
            f"任务书显示，{task['title']}是项目中的独立课题之一，围绕该课题设置了专门研究内容、预期目标、考核指标以及承担单位分工，"
            "体现出从理论研究到工程应用验证的完整实施要求。"
        )
    return [task_map[key] for key in sorted(task_map, key=lambda x: int(x))]


def build_toc(tasks: list[dict[str, object]]) -> list[tuple[str, str, str]]:
    toc = [
        ("TOC 1", "Heading 1", "1 前言"),
        ("TOC 2", "Heading 2", "1.1 研究背景"),
        ("TOC 2", "Heading 2", "1.2 研究现状"),
    ]
    for idx, task in enumerate(tasks, start=1):
        toc.append(("TOC 3", "Heading 3", f"1.2.{idx} {intro_task_heading(task['title'])}"))
    toc.append(("TOC 2", "Heading 2", "1.3 研究内容"))
    for idx, task in enumerate(tasks, start=1):
        toc.append(("TOC 3", "Heading 3", f"1.3.{idx} {shorten_heading_sentence(intro_task_heading(task['title']))}"))
    toc.append(("TOC 2", "Heading 2", "1.4 技术路线"))
    for idx, task in enumerate(tasks, start=1):
        toc.append(("TOC 3", "Heading 3", f"1.4.{idx} {shorten_heading_sentence(intro_task_heading(task['title']))}"))
    for idx, task in enumerate(tasks, start=2):
        toc.append(("TOC 1", "Heading 1", f"{idx} {task['title']}"))
        toc.append(("TOC 2", "Heading 2", f"{idx}.1 概述"))
        for sub_idx, item in enumerate(task["content_items"], start=2):
            toc.append(("TOC 2", "Heading 2", f"{idx}.{sub_idx} {item}"))
            for point_idx, point in enumerate(task.get("specific_points_map", {}).get(item, []), start=1):
                toc.append(("TOC 3", "Heading 3", f"{idx}.{sub_idx}.{point_idx} {shorten_heading_sentence(point)}"))
        summary_no = len(task["content_items"]) + 2
        toc.append(("TOC 2", "Heading 2", f"{idx}.{summary_no} 小结"))
    conclusion_no = len(tasks) + 2
    toc.append(("TOC 1", "Heading 1", f"{conclusion_no} 结论"))
    toc.append(("TOC 1", "Heading 1", "参考文献"))
    return toc


def build_body(tasks: list[dict[str, object]], conclusion_no: int, source_bundle: dict[str, object]) -> dict[str, list[str]]:
    fields = source_bundle["fields"]
    task_titles = "、".join(task["title"] for task in tasks)
    project_name = fields.get("项目名称", "项目")
    business_context = project_business_context(project_name, tasks)
    body = {
        "1.1 研究背景": [
            expand_to_min_length(
                f"《{project_name}》任务书表明，本项目面向{business_context}等典型业务需求展开研究，重点解决复杂电网场景下人工智能分析能力、专家经验利用效率和人机协同决策机制不足的问题。",
                context=project_name,
            ),
            expand_to_min_length(
                f"从任务书设置的课题结构看，项目围绕应用框架与协作机制、可解释机器学习方法、知识引导学习的风险管控决策技术以及决策模块研发与应用验证等方向形成了递进式研究布局，目标是构建既具备算法效率又具备业务可信性的混合增强智能技术体系。",
                context=project_name,
            ),
        ],
        "1.2 研究现状": [
            expand_to_min_length(
                f"结合任务书内容可以看出，当前面向{business_context}的智能分析与决策方法仍普遍存在模型决策透明度不足、人工经验难以结构化融入、复杂工况下样本获取成本高以及在线应用场景中人机协作链路不清晰等问题。",
                context=project_name,
            ),
            expand_to_min_length(
                "因此，需要以电网业务机理、运行规程和风险约束为基础，把数据驱动方法、知识引导机制、可解释分析能力和人工交互能力有机结合起来，形成能够服务实际调控与运检业务的专业化技术报告和实现方案。",
                context=project_name,
            ),
        ],
        "1.3 研究内容": [
            expand_to_min_length(
                f"依据任务书，项目总体研究内容包括：{task_titles}。各课题分别从架构机制、模型解释、知识引导和工程验证等角度展开，覆盖了从基础理论研究到业务落地应用的完整技术链条。",
                context=project_name,
            ),
            expand_to_min_length(
                f"这些研究内容最终服务于地区电网可靠性分析、薄弱环节定位、风险预警、热稳安全校正控制和转供策略生成等关键场景，强调在真实业务约束下实现模型可用、结果可信、交互有效和系统可部署。",
                context=project_name,
            ),
        ],
        "1.4 技术路线": [
            expand_to_min_length(
                "结合任务书内容，项目技术路线可以概括为业务场景分析与问题定义、人机混合增强智能架构设计、模型解释与知识嵌入方法研究、决策模块研发集成以及典型场景应用验证五个层次。",
                context=project_name,
            ),
            expand_to_min_length(
                "任务书中的预期目标和考核指标不仅约束论文、专利和软件成果，更明确了精度提升、识别准确率、节点规模支撑和在线分析时效等工程化指标，因此技术路线必须贯穿方法研究、系统实现与场景验证全过程。",
                context=project_name,
            ),
        ],
    }
    for idx, task in enumerate(tasks, start=1):
        short = intro_task_heading(task["title"])
        body[f"1.2.{idx} {short}"] = task_status_summary(task)
        body[f"1.3.{idx} {shorten_heading_sentence(short)}"] = task_content_summary(task)
        body[f"1.4.{idx} {shorten_heading_sentence(short)}"] = task_route_summary(task)
    for idx, task in enumerate(tasks, start=2):
        prefix = f"{idx}"
        body[f"{prefix}.1 概述"] = task_overview_paragraphs(task)
        for sub_idx, item in enumerate(task["content_items"], start=2):
            key = f"{prefix}.{sub_idx} {item}"
            specific_points = task.get("specific_points_map", {}).get(item, [])
            body[key] = item_section_paragraphs(task, item)
            for point_idx, point in enumerate(specific_points, start=1):
                point_key = f"{prefix}.{sub_idx}.{point_idx} {shorten_heading_sentence(point)}"
                body[point_key] = point_implementation_paragraphs(task, item, point)
        summary_no = len(task["content_items"]) + 2
        body[f"{prefix}.{summary_no} 小结"] = [expand_to_min_length(summarize_task(task), context=task["title"])]
    body[f"{conclusion_no} 结论"] = []
    for idx, task in enumerate(tasks, start=2):
        body[f"{conclusion_no} 结论"].append(
            expand_to_min_length(
                f"课题{idx - 1}“{task['title']}”围绕任务书设定的研究目标，形成了从关键方法研究、技术路线设计到工程验证准备的较完整工作链条。",
                context=task["title"],
            )
        )
        if task.get("content_items"):
            joined = "、".join(task["content_items"][:3])
            body[f"{conclusion_no} 结论"].append(
                expand_to_min_length(
                    f"其中，{joined}等研究内容共同支撑了该课题的核心技术框架，为后续成果凝练、系统研发和应用落地提供了直接依据。",
                    context=task["title"],
                )
            )
        body[f"{conclusion_no} 结论"].append(
            expand_to_min_length(
                "从任务分工和实施安排来看，该课题与项目整体技术体系衔接紧密，能够为企业级电网智能计算推演能力建设提供对应的理论、方法或平台支撑。",
                context=task["title"],
            )
        )
    body["参考文献"] = [
        expand_to_min_length(
            "本文文字内容依据《公司科技项目任务书》整理撰写。由于输入材料未提供参考文献条目，本节不再引用模板原有参考文献内容。",
            context=fields.get("项目名称", "项目"),
        )
    ]
    body_override = source_bundle.get("body_overrides", {})
    if isinstance(body_override, dict):
        for heading, paragraphs in body_override.items():
            if isinstance(heading, str) and isinstance(paragraphs, list):
                body[heading] = [str(p).strip() for p in paragraphs if str(p).strip()]
    return body


def render_pdf(docx_path: Path, pdf_path: Path) -> None:
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    profile = f"file://{tempfile.gettempdir()}/lo_profile_codex_report"
    subprocess.run(
        [
            "soffice",
            f"-env:UserInstallation={profile}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_path.parent),
            str(docx_path),
        ],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )


def pdf_page_text(pdf_path: Path, page_no: int) -> str:
    result = subprocess.run(
        ["pdftotext", "-f", str(page_no), "-l", str(page_no), "-layout", str(pdf_path), "-"],
        check=True,
        capture_output=True,
        text=True,
    )
    return normalize(result.stdout)


def detect_body_pdf_start(pdf_path: Path) -> int:
    all_pages = int(
        re.search(
            r"Pages:\s+(\d+)",
            subprocess.run(["pdfinfo", str(pdf_path)], check=True, capture_output=True, text=True).stdout,
        ).group(1)
    )
    for page in range(1, all_pages + 1):
        text = pdf_page_text(pdf_path, page)
        if normalize("1 前言") in text and normalize("目 录") not in text:
            return page
    return 1


def heading_page_map(pdf_path: Path, toc: list[tuple[str, str, str]], body_pdf_start: int | None = 8) -> dict[str, int]:
    page_map: dict[str, int] = {}
    all_pages = int(
        re.search(
            r"Pages:\s+(\d+)",
            subprocess.run(["pdfinfo", str(pdf_path)], check=True, capture_output=True, text=True).stdout,
        ).group(1)
    )
    if body_pdf_start is None:
        body_pdf_start = detect_body_pdf_start(pdf_path)
    page_texts = {page: pdf_page_text(pdf_path, page) for page in range(body_pdf_start, all_pages + 1)}
    for _, _, heading in toc:
        needle = normalize(heading)
        m = re.match(r"^(\d+(?:\.\d+)*)", heading)
        number_prefix = m.group(1) if m else ""
        heading_prefix = normalize(re.sub(r"[；;（(].*$", "", heading))
        variants = [needle, heading_prefix]
        shortened = heading
        for suffix in ["中的应用验证", "技术研究", "方法研究", "机制研究", "方案设计"]:
            if suffix in shortened:
                shortened = shortened.rsplit(suffix, 1)[0]
                variants.append(normalize(shortened))
        for page, text in page_texts.items():
            if any(variant and len(variant) >= 8 and variant in text for variant in variants):
                page_map[heading] = page - body_pdf_start + 1
                break
            if number_prefix and re.search(rf"(^|\\s){re.escape(number_prefix)}(\\s|\\.)", text):
                page_map[heading] = page - body_pdf_start + 1
                break
    for _, _, heading in toc:
        if heading in page_map:
            continue
        m = re.match(r"^(\d+\.\d+)\s", heading)
        if m:
            prefix = f"{m.group(1)}."
            child_pages = [
                page
                for _, _, candidate in toc
                if candidate in page_map and candidate.startswith(prefix)
            ]
            if child_pages:
                page_map[heading] = min(child_pages)
                continue
        m = re.match(r"^(\d+)\s", heading)
        if m:
            chapter_no = m.group(1)
            sibling = f"{chapter_no}.1 概述"
            if chapter_no == "1":
                sibling = "1.1 研究背景"
            if sibling in page_map:
                page_map[heading] = page_map[sibling]
    for toc_style, _, heading in toc:
        if toc_style != "TOC 2":
            continue
        m = re.match(r"^(\d+\.\d+)\s", heading)
        if not m:
            continue
        prefix = f"{m.group(1)}."
        child_pages = sorted(
            {
                page_map[candidate]
                for child_style, _, candidate in toc
                if child_style == "TOC 3" and candidate in page_map and candidate.startswith(prefix)
            }
        )
        if child_pages:
            page_map[heading] = child_pages[0]
            continue
        chapter_prefix = f"{m.group(1).split('.')[0]}."
        sibling_pages = sorted(
            {
                page_map[candidate]
                for child_style, _, candidate in toc
                if child_style == "TOC 2"
                and candidate in page_map
                and candidate != heading
                and candidate.startswith(chapter_prefix)
            }
        )
        if sibling_pages:
            if heading.endswith("小结"):
                page_map[heading] = sibling_pages[-1]
            else:
                page_map[heading] = sibling_pages[0]
    for toc_style, _, heading in toc:
        if toc_style != "TOC 1" or heading in page_map:
            continue
        m = re.match(r"^(\d+)\s", heading)
        if not m:
            continue
        chapter_prefix = f"{m.group(1)}."
        sibling_pages = sorted(
            {
                page_map[candidate]
                for child_style, _, candidate in toc
                if child_style == "TOC 2"
                and candidate in page_map
                and candidate.startswith(chapter_prefix)
            }
        )
        if sibling_pages:
            page_map[heading] = sibling_pages[0]
    return page_map


def apply_toc_paragraph_format(paragraph, level: int) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.left_indent = Pt(18 * (level - 1)) if level > 1 else None
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    tabs = paragraph.paragraph_format.tab_stops
    tabs.add_tab_stop(Pt(445), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)


def toc_display_text(heading: str, level: int) -> str:
    if level == 2:
        heading = heading.replace("面向企业级电网智能计算推演模型的", "")
        heading = heading.replace("面向企业级电网计算推演的", "")
        heading = heading.replace("企业级电网智能计算推演服务", "计算推演服务")
        heading = heading.replace("企业级电网高维异构图", "高维异构图")
        heading = heading.replace("基于机器学习自适应建模的", "")
        heading = heading.replace("企业级电网智能计算推演共性服务研发及", "共性服务研发及")
        heading = heading.replace("多源数据接入共享与共性服务集成技术框架", "数据接入与集成框架")
        heading = heading.replace("训练、迭代及更新服务架构", "训练迭代更新架构")
        heading = heading.replace("计算推演服务并行计算及其自适应调优技术", "并行计算与自适应调优技术")
        heading = heading.replace("高维异构图模型一体化智能计算技术研究", "高维异构图一体化智能计算")
        heading = heading.replace("共性服务研发及试点应用验证", "共性服务研发与试点验证")
        heading = heading.replace("基于计算范式融合的电网状态一体化智能计算技术", "计算范式融合智能计算")
        heading = heading.replace("基于高通量图计算的电网一体化计算加速技术", "高通量图计算加速技术")
        heading = heading.replace("多时空尺度预测推演类场景的自动化建模技术", "多时空自动化建模技术")
        heading = heading.replace("预测模型的增量式更新和快速迭代优化技术", "预测模型增量更新优化技术")
        heading = heading.replace("面向推演结果校验的模型可解释溯因与反馈更新技术", "模型溯因与反馈更新技术")
    if level == 3:
        heading = re.sub(r"[；;（(].*$", "", heading)
        heading = heading.replace("综合来看，", "")
        heading = heading.replace("围绕核心研究目标形成了较完整的技术链条", "技术链条完整")
        heading = heading.replace("既覆盖关键方法和实现路径，也兼顾工程适配与应用验证", "兼顾方法实现与工程验证")
        heading = heading.replace("能够为课题后续成果凝练和工程落地提供支撑", "支撑成果凝练与工程落地")
    if level == 1:
        heading = heading.replace("企业级电网智能计算推演共性服务技术体系研究", "计算推演共性服务技术体系研究")
        heading = heading.replace("电网高维异构图模型一体化智能计算技术研究", "高维异构图一体化智能计算技术研究")
    return heading


def add_toc_entry(paragraph, heading: str, level: int, page_no: int) -> None:
    clear_paragraph_content(paragraph)
    apply_toc_paragraph_format(paragraph, level)
    display_text = toc_display_text(heading, level)
    add_hyperlink_run(paragraph, f"bm_{heading}", display_text)
    run = paragraph.add_run(f"\t{page_no}")
    run.font.name = "宋体"
    run.font.size = Pt(12)
    set_rfonts(run, ascii_font="宋体", hansi_font="宋体", east_asia_font="宋体", cs_font="宋体")


def replace_manual_toc(docx_path: Path, toc: list[tuple[str, str, str]], page_map: dict[str, int]) -> None:
    doc = Document(str(docx_path))
    title_idx = next(i for i, p in enumerate(doc.paragraphs) if normalize(p.text) == "目 录")
    first_heading_idx = next(i for i, p in enumerate(doc.paragraphs) if normalize(p.text) == "1 前言")
    toc_paras = doc.paragraphs[title_idx + 1:first_heading_idx]
    anchor = None
    for para in toc_paras:
        if para.style.name.startswith("TOC"):
            anchor = para
            break
    if anchor is None:
        anchor = doc.paragraphs[title_idx + 1]
    for para in toc_paras:
        if para is anchor:
            continue
        ppr = para._p.pPr
        if ppr is not None and ppr.sectPr is not None:
            continue
        para._element.getparent().remove(para._element)
    last = anchor
    first = True
    for toc_style, _, heading in toc:
        level = 1 if toc_style == "TOC 1" else 2 if toc_style == "TOC 2" else 3
        page_no = page_map.get(heading, "")
        target = last if first else insert_paragraph_after(last, toc_style)
        target.style = toc_style
        add_toc_entry(target, heading, level, page_no)
        last = target
        first = False
    first_heading = next((p for p in doc.paragraphs if normalize(p.text) == "1 前言"), None)
    if first_heading is not None:
        prev = first_heading._p.getprevious()
        has_page_break = False
        while prev is not None:
            xml = prev.xml
            if "w:type=\"page\"" in xml:
                has_page_break = True
                break
            if "w:t" in xml:
                break
            prev = prev.getprevious()
        if not has_page_break:
            break_para = insert_paragraph_after(last)
            break_para.add_run().add_break(WD_BREAK.PAGE)
    doc.save(docx_path)


def make_summary_from_bundle(source_bundle: dict[str, object], tasks: list[dict[str, object]]) -> tuple[str, str, str, str]:
    front_override = source_bundle.get("front_matter", {})
    if isinstance(front_override, dict):
        summary = front_override.get("summary")
        keywords = front_override.get("keywords")
        abstract = front_override.get("abstract")
        en_keywords = front_override.get("en_keywords")
        if all(isinstance(x, str) and x.strip() for x in [summary, keywords, abstract, en_keywords]):
            return summary.strip(), keywords.strip(), abstract.strip(), en_keywords.strip()
    fields = source_bundle["fields"]
    summary = (
        f"本技术报告依据《公司科技项目任务书》编制，围绕“{fields.get('项目名称', '')}”项目展开。"
        f"任务书明确项目由{fields.get('主要承担单位', '')}牵头实施，研究周期为{fields.get('起止时间', '')}，"
        f"联合多家单位围绕企业级电网智能计算推演共性服务、高维异构图一体化智能计算、多时空推演以及共性服务研发与试点应用验证等方向开展协同攻关。"
        f"报告在保持《纵向科技项目技术报告编写模板-2023版（新模板）》版式要求的基础上，按照任务书中的{len(tasks)}个课题进行展开，"
        "重点说明各课题的研究内容、预期目标、考核指标以及承担单位与经费安排，以形成与项目实施要求一致的技术报告正文。"
    )
    keywords = "电网智能计算推演；共性服务；高维异构图；多时空推演；自适应建模"
    abstract = (
        "This report is prepared from the project assignment document for the study on key common service technologies for enterprise-level intelligent computation and simulation of power grids. "
        "Following the fixed technical report template, the report reorganizes the source material into a formal report structure and describes the background, key research tasks, expected objectives, assessment indicators, and implementation arrangements of each task."
    )
    en_keywords = "intelligent computation and simulation; common service; heterogeneous graph; multi-temporal simulation; adaptive modeling"
    return summary, keywords, abstract, en_keywords


def add_front_matter(doc: Document, samples: dict[str, object], source_bundle: dict[str, object], tasks: list[dict[str, object]]) -> None:
    summary, keywords, abstract, en_keywords = make_summary_from_bundle(source_bundle, tasks)
    add_paragraph(doc, "内容摘要", "Normal", samples["summary_title"], "黑体")
    blank = doc.add_paragraph(style="Normal")
    blank.paragraph_format.first_line_indent = None
    add_paragraph(doc, summary, "Normal", samples["body"], "宋体")
    add_paragraph(doc, f"关键词：{keywords}", "Normal", samples["body"], "宋体")
    add_page_break(doc)
    p = add_paragraph(doc, "ABSTRACT", "Normal", samples["summary_title"], "黑体")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(18)
        run.bold = True
        set_rfonts(run, ascii_font="Times New Roman", hansi_font="Times New Roman", east_asia_font="Times New Roman", cs_font="Times New Roman")
    blank = doc.add_paragraph(style="Normal")
    blank.paragraph_format.first_line_indent = None
    p = add_paragraph(doc, abstract, "Normal", samples["body"], "宋体")
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        set_rfonts(run, ascii_font="Times New Roman", hansi_font="Times New Roman", east_asia_font="Times New Roman", cs_font="Times New Roman")
    p = add_paragraph(doc, f"KEYWORDS: {en_keywords}", "Normal", samples["body"], "宋体")
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        set_rfonts(run, ascii_font="Times New Roman", hansi_font="Times New Roman", east_asia_font="Times New Roman", cs_font="Times New Roman")
    add_page_break(doc)
    p = add_paragraph(doc, "目 录", "Normal", samples["toc_title"], "黑体")
    p.alignment = samples["toc_title"].alignment
    blank = doc.add_paragraph(style="Normal")
    blank.paragraph_format.first_line_indent = None
    p = doc.add_paragraph(style="TOC 1")
    p.paragraph_format.first_line_indent = None
    add_toc_field(p)


def add_body(doc: Document, samples: dict[str, object], toc: list[tuple[str, str, str]], body_map: dict[str, list[str]]) -> None:
    bookmark_id = 1
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    for _, style_name, heading in toc:
        if style_name == "Heading 1" and heading != "1 前言":
            add_page_break(doc)
        sample = samples["h1"] if style_name == "Heading 1" else samples["h2"] if style_name == "Heading 2" else samples["h3"]
        p = add_paragraph(doc, heading, style_name, sample, "黑体", copy_format=False)
        if p._p.pPr is not None and p._p.pPr.numPr is not None:
            p._p.pPr.remove(p._p.pPr.numPr)
        add_bookmark(p, bookmark_id, f"bm_{heading}")
        bookmark_id += 1
        for para in body_map.get(heading, []):
            add_paragraph(doc, para, "Normal", samples["body"], "宋体")


def normalize_heading_paragraphs(doc: Document) -> None:
    heading_rules = {
        "Heading 1": {"font": "黑体", "size": Pt(15), "bold": True},
        "Heading 2": {"font": "黑体", "size": Pt(14), "bold": True},
        "Heading 3": {"font": "黑体", "size": Pt(12), "bold": False},
    }
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name not in heading_rules:
            continue
        rule = heading_rules[style_name]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = paragraph.paragraph_format
        pf.first_line_indent = None
        pf.left_indent = None
        pf.right_indent = None
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pf.line_spacing = 1.5
        for run in paragraph.runs:
            run.font.name = rule["font"]
            run.font.size = rule["size"]
            run.bold = rule["bold"]
            set_rfonts(run, ascii_font=rule["font"], hansi_font=rule["font"], east_asia_font=rule["font"], cs_font=rule["font"])


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a fixed-format technical report DOCX.")
    parser.add_argument(
        "--template",
        default=str(DEFAULT_TEMPLATE),
        help="Path to the fixed technical-report template DOCX.",
    )
    parser.add_argument(
        "--source",
        default="8_195_山东公司_面向企业级电网智能计算推演的共性服务关键技术研究_任务书（下达稿）.pdf",
        help="Path to the source content document, typically a PDF or DOCX.",
    )
    parser.add_argument(
        "--output",
        default="technical-report-final-format.docx",
        help="Path to the generated output DOCX.",
    )
    parser.add_argument(
        "--front-reference",
        default=str(DEFAULT_FRONT_REFERENCE),
        help="Optional DOCX whose front matter before 内容摘要 should be copied as the fixed cover/sign-off layout.",
    )
    parser.add_argument(
        "--english-title",
        default="",
        help="Optional English project title supplied by the calling LLM for the cover page.",
    )
    parser.add_argument(
        "--llm-overrides",
        default="",
        help="Optional JSON file containing LLM-extracted task data, front matter, and body overrides.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    template_path = Path(args.template)
    source_path = Path(args.source)
    out = Path(args.output)
    front_reference = Path(args.front_reference) if args.front_reference else None
    llm_overrides_path = Path(args.llm_overrides) if args.llm_overrides else None
    if not template_path.is_absolute():
        template_path = (Path.cwd() / template_path).resolve()
    if not source_path.is_absolute():
        source_path = (Path.cwd() / source_path).resolve()
    if not out.is_absolute():
        out = (Path.cwd() / out).resolve()
    if front_reference and not front_reference.is_absolute():
        front_reference = (Path.cwd() / front_reference).resolve()
    if llm_overrides_path and not llm_overrides_path.is_absolute():
        llm_overrides_path = (Path.cwd() / llm_overrides_path).resolve()
    template = Document(str(template_path))
    output_base = front_reference if front_reference and front_reference.exists() else template_path
    output = Document(str(output_base))
    source_bundle = parse_pdf_source(source_path)
    source_bundle = apply_llm_overrides(source_bundle, load_llm_overrides(llm_overrides_path))
    if args.english_title:
        source_bundle["fields"]["英文标题"] = normalize(args.english_title)

    tasks = source_bundle["tasks"]
    toc = build_toc(tasks)
    conclusion_no = len(tasks) + 2
    body_map = build_body(tasks, conclusion_no, source_bundle)

    samples = {
        "summary_title": find_template_paragraph(template, exact_text="内容摘要"),
        "toc_title": find_template_paragraph(template, exact_text="目 录"),
        "h1": find_template_paragraph(template, style_name="Heading 1"),
        "h2": find_template_paragraph(template, style_name="Heading 2"),
        "h3": find_template_paragraph(template, style_name="3级"),
        "body": find_template_paragraph(template, startswith="正文宋体小四"),
    }

    ensure_report_styles(output, template)
    apply_document_layout(output)
    set_update_fields_on_open(output)

    if front_reference and front_reference.exists():
        replace_fixed_front_matter_fields(output, source_bundle)
    else:
        replace_cover_from_bundle(output, source_bundle)
        postprocess_cover_layout(output)
    clear_body_after(output, "内容摘要")
    add_front_matter(output, samples, source_bundle, tasks)
    add_body(output, samples, toc, body_map)
    normalize_heading_paragraphs(output)
    apply_document_layout(output)
    if len(output.sections) >= 3:
        summary_section = output.sections[-2]
        body_section = output.sections[-1]
        configure_footer_page_number(summary_section)
        set_section_page_number(summary_section, fmt="lowerRoman", start=1)
        configure_footer_page_number(body_section)
        set_section_page_number(body_section, fmt="decimal", start=1)

    output.save(out)
    pdf_dir = Path(tempfile.gettempdir()) / "codex_report_toc"
    pdf_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = pdf_dir / f"{out.stem}.pdf"
    render_pdf(out, pdf_path)
    body_pdf_start = detect_body_pdf_start(pdf_path)
    page_map = heading_page_map(pdf_path, toc, body_pdf_start=body_pdf_start)
    replace_manual_toc(out, toc, page_map)
    try:
        pdf_path.unlink(missing_ok=True)
    except OSError:
        pass
    print(out)


if __name__ == "__main__":
    main()
