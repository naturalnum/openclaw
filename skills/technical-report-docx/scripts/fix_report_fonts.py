#!/usr/bin/env python3
"""Force generated report paragraphs to match template font and paragraph format."""

from __future__ import annotations

import argparse
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn


def normalize(text: str) -> str:
    return " ".join(text.split())


def copy_paragraph_format(dst, src) -> None:
    dst.alignment = src.alignment
    dpf = dst.paragraph_format
    spf = src.paragraph_format
    dpf.left_indent = spf.left_indent
    dpf.right_indent = spf.right_indent
    dpf.first_line_indent = spf.first_line_indent
    dpf.keep_together = spf.keep_together
    dpf.keep_with_next = spf.keep_with_next
    dpf.page_break_before = spf.page_break_before
    dpf.widow_control = spf.widow_control
    dpf.space_before = spf.space_before
    dpf.space_after = spf.space_after
    dpf.line_spacing = spf.line_spacing
    dpf.line_spacing_rule = spf.line_spacing_rule


def remove_numpr(paragraph) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.numPr
    if numPr is not None:
        pPr.remove(numPr)


def remove_style_numpr(style) -> None:
    pPr = style.element.get_or_add_pPr()
    numPr = pPr.numPr
    if numPr is not None:
        pPr.remove(numPr)


def ensure_rfonts(run, ascii_font=None, hansi_font=None, east_asia_font=None, cs_font=None) -> None:
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = deepcopy(run._element._new_rPr().rFonts)  # pragma: no cover
    if rFonts is None:
        from docx.oxml import OxmlElement

        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    if ascii_font:
        rFonts.set(qn("w:ascii"), ascii_font)
    if hansi_font:
        rFonts.set(qn("w:hAnsi"), hansi_font)
    if east_asia_font:
        rFonts.set(qn("w:eastAsia"), east_asia_font)
    if cs_font:
        rFonts.set(qn("w:cs"), cs_font)


def copy_run_format(dst_run, src_run) -> None:
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    dst_run.font.size = src_run.font.size
    dst_run.font.name = src_run.font.name
    if src_run.font.color and src_run.font.color.rgb:
        dst_run.font.color.rgb = src_run.font.color.rgb

    src_rpr = src_run._r.rPr
    if src_rpr is None or src_rpr.rFonts is None:
        return
    rfonts = src_rpr.rFonts
    ensure_rfonts(
        dst_run,
        ascii_font=rfonts.get(qn("w:ascii")),
        hansi_font=rfonts.get(qn("w:hAnsi")),
        east_asia_font=rfonts.get(qn("w:eastAsia")),
        cs_font=rfonts.get(qn("w:cs")),
    )


def find_paragraph(doc: Document, text: str, occurrence: str = "first") -> int:
    matches = []
    for i, p in enumerate(doc.paragraphs):
        if normalize(p.text) == text:
            matches.append(i)
    if matches:
        return matches[-1] if occurrence == "last" else matches[0]
    raise ValueError(f"paragraph not found: {text}")


def first_nonempty_run(paragraph):
    for run in paragraph.runs:
        if run.text.strip():
            return run
    return paragraph.runs[0] if paragraph.runs else None


def apply_heading_format(paragraph, sample) -> None:
    paragraph.style = sample.style
    copy_paragraph_format(paragraph, sample)
    remove_numpr(paragraph)
    sample_run = first_nonempty_run(sample)
    if sample_run is None:
        return
    if not paragraph.runs:
        paragraph.add_run(paragraph.text)
        paragraph.text = ""
    for run in paragraph.runs:
        copy_run_format(run, sample_run)


def apply_body_format(paragraph, sample) -> None:
    paragraph.style = sample.style
    copy_paragraph_format(paragraph, sample)
    remove_numpr(paragraph)
    sample_run = first_nonempty_run(sample)
    if sample_run is None:
        return
    for run in paragraph.runs:
        copy_run_format(run, sample_run)
        text = run.text
        has_ascii = any(ord(ch) < 128 for ch in text if ch.strip())
        if has_ascii:
            ensure_rfonts(
                run,
                ascii_font="Times New Roman",
                hansi_font="Times New Roman",
                east_asia_font="宋体",
                cs_font="Times New Roman",
            )
        else:
            ensure_rfonts(
                run,
                ascii_font="Times New Roman",
                hansi_font="Times New Roman",
                east_asia_font="宋体",
                cs_font="Times New Roman",
            )


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True)
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    template = Document(args.template)
    doc = Document(args.input)

    for style_name in ["Heading 1", "Heading 3"]:
        if style_name in [s.name for s in doc.styles]:
            remove_style_numpr(doc.styles[style_name])

    chapter_sample = template.paragraphs[115]
    section_sample = template.paragraphs[116]
    sub_section_sample = template.paragraphs[167]
    body_sample = template.paragraphs[284]
    reference_heading = template.paragraphs[1683]
    reference_body = template.paragraphs[1684]

    body_start = find_paragraph(doc, "1 前言", occurrence="last")
    ref_start = find_paragraph(doc, "参考文献", occurrence="last")

    for i, paragraph in enumerate(doc.paragraphs[body_start:], start=body_start):
        text = normalize(paragraph.text)
        if not text:
            continue
        if i == ref_start:
            apply_heading_format(paragraph, reference_heading)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            continue
        if text.startswith("参考文献"):
            apply_heading_format(paragraph, reference_heading)
            continue
        if paragraph.style.name in {"Heading 1"}:
            apply_heading_format(paragraph, chapter_sample)
            paragraph.alignment = None
            continue
        if paragraph.style.name in {"002二级标题"}:
            apply_heading_format(paragraph, section_sample)
            paragraph.alignment = None
            continue
        if paragraph.style.name in {"003三级标题", "Heading 3"}:
            apply_heading_format(paragraph, sub_section_sample)
            paragraph.alignment = None
            continue
        if i > ref_start:
            apply_body_format(paragraph, reference_body)
        else:
            apply_body_format(paragraph, body_sample)
        paragraph.alignment = None

    doc.save(args.output)


if __name__ == "__main__":
    main()
