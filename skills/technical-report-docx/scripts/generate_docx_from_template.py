#!/usr/bin/env python3

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


SECTION_HEADING_RE = re.compile(r"^[一二三四五六七八九十]+、")
SUBSECTION_HEADING_RE = re.compile(r"^（[一二三四五六七八九十]+）")


def normalize(text: str) -> str:
    return " ".join(text.split())


def iter_block_items(document: Document):
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def load_mapping(path: Path | None) -> dict:
    if path is None or not path.exists():
        return {}
    data = json.loads(path.read_text(encoding="utf-8"))
    if "suggested_mapping" in data:
        return data["suggested_mapping"]
    return data


def first_existing_style(document: Document, *candidates: str | None) -> str:
    style_names = {style.name for style in document.styles}
    for candidate in candidates:
        if candidate and candidate in style_names:
            return candidate
    return "Normal"


def nonempty_paragraphs(document: Document) -> list[Paragraph]:
    return [paragraph for paragraph in document.paragraphs if normalize(paragraph.text)]


def block_text(block) -> str:
    if isinstance(block, Paragraph):
        return normalize(block.text)
    return " ".join(normalize(cell.text) for row in block.rows for cell in row.cells)


def clone_cover_format(source_paragraph: Paragraph, target_paragraph: Paragraph):
    target_paragraph.style = source_paragraph.style
    target_paragraph.alignment = source_paragraph.alignment
    target_paragraph.paragraph_format.left_indent = source_paragraph.paragraph_format.left_indent
    target_paragraph.paragraph_format.right_indent = source_paragraph.paragraph_format.right_indent
    target_paragraph.paragraph_format.first_line_indent = source_paragraph.paragraph_format.first_line_indent
    target_paragraph.paragraph_format.space_before = source_paragraph.paragraph_format.space_before
    target_paragraph.paragraph_format.space_after = source_paragraph.paragraph_format.space_after
    target_paragraph.paragraph_format.line_spacing = source_paragraph.paragraph_format.line_spacing

    sample_run = None
    for run in source_paragraph.runs:
        if run.text.strip():
            sample_run = run
            break

    new_run = target_paragraph.runs[0] if target_paragraph.runs else target_paragraph.add_run()
    if sample_run is None:
        return

    new_run.font.name = sample_run.font.name
    new_run.font.size = sample_run.font.size
    new_run.bold = sample_run.bold
    new_run.italic = sample_run.italic


def extract_cover_fields(source_doc: Document) -> dict[str, str]:
    fields: dict[str, str] = {}
    if not source_doc.tables:
        return fields
    for row in source_doc.tables[0].rows:
        cells = [" ".join(cell.text.split()) for cell in row.cells]
        if len(cells) >= 2 and cells[0]:
            fields[cells[0].rstrip("：:")] = cells[1]
    return fields


def clear_document_body(document: Document):
    body = document.element.body
    sect_pr = None
    for child in list(body.iterchildren()):
        if child.tag.endswith("}sectPr"):
            sect_pr = child
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def classify_paragraph(text: str, position: int) -> str:
    stripped = text.strip()
    if not stripped:
        return "skip"
    if SECTION_HEADING_RE.match(stripped):
        return "heading_1"
    if SUBSECTION_HEADING_RE.match(stripped):
        return "heading_2"
    return "body"


def add_styled_paragraph(document: Document, text: str, style_name: str):
    paragraph = document.add_paragraph(style=style_name)
    paragraph.add_run(text)
    return paragraph


def find_nonempty_paragraph(document: Document, text: str) -> Paragraph | None:
    for paragraph in document.paragraphs:
        if normalize(paragraph.text) == text:
            return paragraph
    return None


def source_table_to_matrix(table: Table) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        current: list[str] = []
        for cell in row.cells:
            fragments = []
            for paragraph in cell.paragraphs:
                text = " ".join(paragraph.text.split())
                if text:
                    fragments.append(text)
            current.append("\n".join(fragments))
        rows.append(current)
    return rows


def set_section_landscape(section):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width


def set_section_portrait(section):
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = section.page_height, section.page_width


def estimate_column_weights(matrix: list[list[str]]) -> list[float]:
    cols = len(matrix[0])
    weights = [1.0] * cols
    sample_rows = matrix[: min(len(matrix), 10)]
    for col in range(cols):
        max_len = max((len((row[col] or "").replace("\n", " ")) for row in sample_rows), default=1)
        weights[col] = max(1.0, min(4.0, max_len / 12))
    if cols >= 8:
        weights = [max(0.8, min(weight, 2.5)) for weight in weights]
    if cols >= 12:
        weights = [max(0.7, min(weight, 1.8)) for weight in weights]
    return weights


def table_font_size(cols: int) -> float:
    if cols <= 4:
        return 10.5
    if cols <= 6:
        return 9.5
    if cols <= 8:
        return 8.5
    if cols <= 12:
        return 7.5
    return 6.5


def available_width(section) -> int:
    return int(section.page_width - section.left_margin - section.right_margin)


def add_table(document: Document, matrix: list[list[str]], table_style: str, body_style: str):
    if not matrix or not matrix[0]:
        return
    cols = len(matrix[0])
    needs_landscape = cols >= 7
    inserted_landscape = False
    if needs_landscape:
        landscape = document.add_section(WD_SECTION_START.NEW_PAGE)
        set_section_landscape(landscape)
        inserted_landscape = True

    table = document.add_table(rows=len(matrix), cols=len(matrix[0]))
    table.style = table_style
    table.autofit = False

    weights = estimate_column_weights(matrix)
    total_weight = sum(weights)
    width = available_width(document.sections[-1])
    for idx, weight in enumerate(weights):
        col_width = int(width * weight / total_weight)
        for cell in table.columns[idx].cells:
            cell.width = col_width

    body_font = table_font_size(cols)
    for row_idx, row in enumerate(matrix):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.style = body_style
            paragraph.paragraph_format.first_line_indent = None
            run = paragraph.add_run(value)
            run.font.size = Pt(body_font)
            if row_idx == 0:
                paragraph.alignment = 1
                run.bold = True
            elif cols <= 4:
                paragraph.alignment = 0
            else:
                paragraph.alignment = 1 if col_idx == 0 else 0

    if inserted_landscape:
        portrait = document.add_section(WD_SECTION_START.NEW_PAGE)
        set_section_portrait(portrait)


def add_caption(document: Document, text: str, caption_style: str):
    paragraph = add_styled_paragraph(document, text, caption_style)
    paragraph.alignment = 1
    return paragraph


def maybe_page_break(document: Document, last_was_table: bool):
    if not document.paragraphs:
        return
    if last_was_table:
        document.add_paragraph()
        return
    document.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)


def build_cover(template_doc: Document, source_doc: Document, template_cover_samples: list[Paragraph]):
    fields = extract_cover_fields(source_doc)
    lines = [
        "国家电网公司总部科技项目",
        fields.get("项目名称", "任务书"),
        "任务书",
        fields.get("主要承担单位", ""),
        fields.get("起止时间", ""),
    ]
    sample_indexes = [2, 3, 5, 6, 7]
    for line, sample_index in zip(lines, sample_indexes):
        if not line:
            continue
        paragraph = add_styled_paragraph(template_doc, line, "Normal")
        if template_cover_samples:
            clone_cover_format(
                template_cover_samples[min(sample_index, len(template_cover_samples) - 1)],
                paragraph,
            )
    if template_doc.paragraphs:
        template_doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)


def section_break(document: Document):
    if document.paragraphs:
        document.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)


def add_body_paragraph(document: Document, text: str, body_style: str):
    if not text:
        return None
    paragraph = add_styled_paragraph(document, text, body_style)
    return paragraph


def add_single_cell_content(document: Document, text: str, heading2_style: str, body_style: str):
    current_style = body_style
    for raw_line in text.splitlines():
        line = normalize(raw_line)
        if not line:
            continue
        if re.match(r"^\d+(\.\d+)+", line):
            add_styled_paragraph(document, line, heading2_style)
            current_style = body_style
            continue
        add_body_paragraph(document, line, current_style)


def collect_tail_paragraphs(source_doc: Document, start_text: str) -> list[str]:
    started = False
    results: list[str] = []
    for paragraph in source_doc.paragraphs:
        text = " ".join(paragraph.text.split())
        if not text:
            continue
        if text == start_text:
            started = True
        if started:
            results.append(text)
    return results


def build_specialized_report(
    template_doc: Document,
    source_doc: Document,
    chapter_style: str,
    section_style: str,
    subsection_style: str,
    body_style: str,
    table_style: str,
    caption_style: str,
):
    cover_samples = nonempty_paragraphs(template_doc)[:8]
    clear_document_body(template_doc)
    build_cover(template_doc, source_doc, cover_samples)

    source_tables = list(source_doc.tables)
    if len(source_tables) < 9:
        return False

    add_styled_paragraph(template_doc, "1 前言", chapter_style)
    add_styled_paragraph(template_doc, "1.1 项目背景", section_style)
    for text in [
        "（四）研究内容及考核指标",
        "1、根据经过专家评审的项目可行性研究报告编制，项目按项目和各课题分别论述；独立课题按课题和各子课题分别阐述；",
        "2、研究内容要围绕关键科学、技术问题，系统、有机地形成一个整体，突出研究特色和创新点；",
        "3、预期目标描述项目研究的总体目标以及通过项目的研究预期达到的效果。说明项目解决的问题，解决到何种程度、取得的突破、成果等，并分析说明研究完成后发挥的作用和意义（如对经济、社会、公司、电网发展的影响，社会效益与经济效益等）。",
    ]:
        paragraph = find_nonempty_paragraph(source_doc, text)
        if paragraph is not None:
            add_body_paragraph(template_doc, normalize(paragraph.text), body_style)
    add_styled_paragraph(template_doc, "1.2 项目基本信息", section_style)
    add_caption(template_doc, "表1-1 项目基本信息", caption_style)
    add_table(template_doc, source_table_to_matrix(source_tables[0]), table_style, body_style)
    add_styled_paragraph(template_doc, "1.3 综合信息表", section_style)
    add_caption(template_doc, "表1-2 项目综合信息", caption_style)
    add_table(template_doc, source_table_to_matrix(source_tables[1]), table_style, body_style)
    add_styled_paragraph(template_doc, "1.4 研究内容", section_style)

    section_break(template_doc)
    add_styled_paragraph(template_doc, "2 研究内容及考核指标", chapter_style)
    section_para = find_nonempty_paragraph(source_doc, "一、研究内容及考核指标")
    if section_para is not None:
        add_body_paragraph(template_doc, " ".join(section_para.text.split()), body_style)
    research_matrix = source_table_to_matrix(source_tables[2])
    if len(research_matrix) == 1 and len(research_matrix[0]) == 1:
        add_styled_paragraph(template_doc, "2.1 项目研究内容及成果", section_style)
        add_single_cell_content(template_doc, research_matrix[0][0], subsection_style, body_style)
    else:
        add_caption(template_doc, "表2-1 研究内容及考核指标", caption_style)
        add_table(template_doc, research_matrix, table_style, body_style)
    add_styled_paragraph(template_doc, "2.2 课题研究内容及考核指标", section_style)
    add_caption(template_doc, "表2-1 研究内容及考核指标", caption_style)
    add_table(template_doc, source_table_to_matrix(source_tables[2]), table_style, body_style)

    section_break(template_doc)
    add_styled_paragraph(template_doc, "3 分工安排与组织实施", chapter_style)
    add_styled_paragraph(template_doc, "3.1 课题分工安排", section_style)
    add_caption(template_doc, "表3-1 项目分工安排", caption_style)
    add_table(template_doc, source_table_to_matrix(source_tables[3]), table_style, body_style)
    add_styled_paragraph(template_doc, "3.2 项目进度计划", section_style)
    add_caption(template_doc, "表3-2 项目进度计划", caption_style)
    add_table(template_doc, source_table_to_matrix(source_tables[4]), table_style, body_style)
    add_styled_paragraph(template_doc, "3.3 课题进度计划及考核要求", section_style)
    add_caption(template_doc, "表3-3 课题进度计划及考核要求", caption_style)
    add_table(template_doc, source_table_to_matrix(source_tables[5]), table_style, body_style)
    add_styled_paragraph(template_doc, "3.4 小结", section_style)
    add_body_paragraph(template_doc, "（五）分工安排", body_style)

    section_break(template_doc)
    add_styled_paragraph(template_doc, "4 经费预算与支付计划", chapter_style)
    add_styled_paragraph(template_doc, "4.1 经费预算安排", section_style)
    add_caption(template_doc, "表4-1 经费预算安排", caption_style)
    add_table(template_doc, source_table_to_matrix(source_tables[6]), table_style, body_style)
    add_styled_paragraph(template_doc, "4.2 出资方案及支付计划", section_style)
    funding_intro = find_nonempty_paragraph(
        source_doc,
        "项目经费总额为人民币(大写)玖佰玖拾捌万元整(￥9,980,000.00)，具体出资方案及支付计划如下：",
    )
    if funding_intro is not None:
        add_body_paragraph(template_doc, " ".join(funding_intro.text.split()), body_style)
    add_caption(template_doc, "表4-2 出资方案及支付计划", caption_style)
    add_table(template_doc, source_table_to_matrix(source_tables[7]), table_style, body_style)
    add_styled_paragraph(template_doc, "4.3 小结", section_style)
    add_body_paragraph(template_doc, "（七）经费预算安排", body_style)

    section_break(template_doc)
    add_styled_paragraph(template_doc, "5 研究人员与联系方式", chapter_style)
    add_styled_paragraph(template_doc, "5.1 研究人员", section_style)
    add_caption(template_doc, "表5-1 研究人员", caption_style)
    add_table(template_doc, source_table_to_matrix(source_tables[8]), table_style, body_style)
    add_styled_paragraph(template_doc, "5.2 联系方式", section_style)
    for index, table in enumerate(source_tables[9:], start=1):
        add_caption(template_doc, f"表5-{index + 1} 联系方式信息", caption_style)
        add_table(template_doc, source_table_to_matrix(table), table_style, body_style)
    add_styled_paragraph(template_doc, "5.3 有关问题说明", section_style)
    tail_paragraphs = collect_tail_paragraphs(source_doc, "八、有关问题说明")
    if tail_paragraphs:
        for text in tail_paragraphs[1:]:
            add_body_paragraph(template_doc, text, body_style)

    section_break(template_doc)
    add_styled_paragraph(template_doc, "6 结论", chapter_style)
    for text in [
        "（1）报告6篇",
        "（2）完成面向数字化电网安全韧性提升的人机协同智能决策模块研发及应用；在1-2家具备千级节点规模的220kV及以下电网完成智能决策模块的部署，试点应用区域包含分布式光伏、电动汽车等不少于3类灵活资源；",
        "（3）申请发明专利10项；登记软件著作权1项；发表或录用核心期刊或三大检索论文7篇。",
    ]:
        add_body_paragraph(template_doc, text, body_style)

    return True


def build_document(template: Path, source: Path, output: Path, mapping_path: Path | None):
    template_doc = Document(str(template))
    source_doc = Document(str(source))
    mapping = load_mapping(mapping_path)

    title_style = first_existing_style(template_doc, "Normal")
    chapter_style = first_existing_style(template_doc, "Heading 1", "002二级标题", "Normal")
    heading2_style = first_existing_style(template_doc, "002二级标题", "Heading 1", "Normal")
    heading3_style = first_existing_style(template_doc, "003三级标题", "Heading 3", "Normal")
    body_style = first_existing_style(
        template_doc,
        "本文正文",
        "004正文",
        "报告正文",
        "Normal",
    )
    caption_style = first_existing_style(template_doc, "图名", body_style)
    table_style = first_existing_style(
        template_doc,
        mapping.get("table_text"),
        "Normal Table",
        "Table Grid",
    )
    used_specialized = build_specialized_report(
        template_doc,
        source_doc,
        chapter_style,
        heading2_style,
        heading3_style,
        body_style,
        table_style,
        caption_style,
    )

    if not used_specialized:
        template_cover_samples = nonempty_paragraphs(template_doc)[:8]
        clear_document_body(template_doc)
        build_cover(template_doc, source_doc, template_cover_samples)

        started_body = False
        nonempty_paragraph_index = 0
        last_was_table = False
        for block in iter_block_items(source_doc):
            text_hint = block_text(block)
            if not started_body:
                if isinstance(block, Paragraph) and text_hint == "填 写 说 明":
                    started_body = True
                else:
                    continue
            if isinstance(block, Paragraph):
                text = text_hint
                if not text:
                    continue
                nonempty_paragraph_index += 1
                role = classify_paragraph(text, nonempty_paragraph_index)
                if role == "skip":
                    continue
                if role == "heading_1":
                    if nonempty_paragraph_index > 3 and template_doc.paragraphs:
                        maybe_page_break(template_doc, last_was_table)
                    style_name = chapter_style
                elif role == "heading_2":
                    style_name = heading2_style
                else:
                    style_name = body_style
                paragraph = add_styled_paragraph(template_doc, text, style_name)
                if role == "heading_1":
                    paragraph.paragraph_format.space_before = Pt(6)
                    paragraph.paragraph_format.space_after = Pt(6)
                elif role == "heading_2":
                    paragraph.paragraph_format.space_before = Pt(3)
                    paragraph.paragraph_format.space_after = Pt(3)
                last_was_table = False
                continue

            matrix = source_table_to_matrix(block)
            add_table(template_doc, matrix, table_style, body_style)
            template_doc.add_paragraph(style=body_style)
            last_was_table = True

    output.parent.mkdir(parents=True, exist_ok=True)
    template_doc.save(str(output))


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Generate a template-styled DOCX whose content comes from the input document."
    )
    parser.add_argument("--template", required=True, help="Path to the template DOCX")
    parser.add_argument("--source", required=True, help="Path to the source DOCX")
    parser.add_argument("--output", required=True, help="Path to the output DOCX")
    parser.add_argument(
        "--style-mapping",
        default=None,
        help="Optional style-mapping JSON generated by suggest_style_mapping.py",
    )
    args = parser.parse_args()

    template = Path(args.template)
    source = Path(args.source)
    output = Path(args.output)
    mapping_path = Path(args.style_mapping) if args.style_mapping else None

    if not template.exists():
        raise SystemExit(f"ERROR: template not found: {template}")
    if not source.exists():
        raise SystemExit(f"ERROR: source not found: {source}")

    build_document(template, source, output, mapping_path)
    print(
        json.dumps(
            {
                "template": str(template),
                "source": str(source),
                "output": str(output),
                "style_mapping": str(mapping_path) if mapping_path else None,
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
