#!/usr/bin/env python3

import argparse
import json
from pathlib import Path

from extract_source_text import extract_text


def run_inspect(template_path: Path):
    try:
        from docx import Document
    except Exception as exc:
        raise SystemExit(
            "ERROR: python-docx is required.\n"
            "Install with:\n"
            "  pip install -r skills/technical-report-docx/scripts/requirements.txt"
        ) from exc

    from collections import Counter

    doc = Document(str(template_path))
    style_counter = Counter()
    font_counter = Counter()
    examples = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style is not None else "(no style)"
        style_counter[style_name] += 1
        for run in para.runs:
            if run.font.name:
                font_counter[run.font.name] += 1
        if len(examples) < 25:
            examples.append((style_name, text[:120]))

    lines = [
        f"# Template Format Summary: {template_path.name}",
        "",
        "## Frequent paragraph styles",
        "",
    ]
    for name, count in style_counter.most_common(20):
        lines.append(f"- {name}: {count}")
    lines.extend(["", "## Frequent explicit fonts", ""])
    for name, count in font_counter.most_common(20):
        lines.append(f"- {name}: {count}")
    lines.extend(["", "## Sample paragraphs", ""])
    for style_name, text in examples:
        lines.append(f"- [{style_name}] {text}")
    return "\n".join(lines)


def pt_value(value):
    if value is None:
        return None
    try:
        return round(float(value.pt), 2)
    except Exception:
        return None


def normalize_alignment(value):
    if value is None:
        return "inherit"
    return str(value)


def run_audit(template_path: Path):
    try:
        from docx import Document
    except Exception as exc:
        raise SystemExit(
            "ERROR: python-docx is required.\n"
            "Install with:\n"
            "  pip install -r skills/technical-report-docx/scripts/requirements.txt"
        ) from exc

    from collections import Counter, defaultdict

    doc = Document(str(template_path))

    global_stats = {
        "paragraph_count": 0,
        "nonempty_paragraph_count": 0,
        "styles": Counter(),
        "fonts": Counter(),
        "sizes_pt": Counter(),
        "alignments": Counter(),
        "bold_true_runs": 0,
        "italic_true_runs": 0,
        "run_count": 0,
    }

    style_profiles = defaultdict(
        lambda: {
            "paragraphs": 0,
            "fonts": Counter(),
            "sizes_pt": Counter(),
            "alignments": Counter(),
            "bold_true_runs": 0,
            "italic_true_runs": 0,
            "run_count": 0,
            "examples": [],
        }
    )

    for para in doc.paragraphs:
        global_stats["paragraph_count"] += 1
        text = para.text.strip()
        if not text:
            continue

        global_stats["nonempty_paragraph_count"] += 1
        style_name = para.style.name if para.style is not None else "(no style)"
        alignment = normalize_alignment(para.alignment)

        global_stats["styles"][style_name] += 1
        global_stats["alignments"][alignment] += 1

        profile = style_profiles[style_name]
        profile["paragraphs"] += 1
        profile["alignments"][alignment] += 1
        if len(profile["examples"]) < 5:
            profile["examples"].append(text[:160])

        for run in para.runs:
            global_stats["run_count"] += 1
            profile["run_count"] += 1

            if run.font.name:
                global_stats["fonts"][run.font.name] += 1
                profile["fonts"][run.font.name] += 1

            size_pt = pt_value(run.font.size)
            if size_pt is not None:
                key = f"{size_pt}pt"
                global_stats["sizes_pt"][key] += 1
                profile["sizes_pt"][key] += 1

            if bool(run.bold):
                global_stats["bold_true_runs"] += 1
                profile["bold_true_runs"] += 1
            if bool(run.italic):
                global_stats["italic_true_runs"] += 1
                profile["italic_true_runs"] += 1

    result = {
        "template": template_path.name,
        "global_stats": {
            "paragraph_count": global_stats["paragraph_count"],
            "nonempty_paragraph_count": global_stats["nonempty_paragraph_count"],
            "run_count": global_stats["run_count"],
            "bold_true_runs": global_stats["bold_true_runs"],
            "italic_true_runs": global_stats["italic_true_runs"],
            "styles": dict(global_stats["styles"].most_common()),
            "fonts": dict(global_stats["fonts"].most_common()),
            "sizes_pt": dict(global_stats["sizes_pt"].most_common()),
            "alignments": dict(global_stats["alignments"].most_common()),
        },
        "style_profiles": {},
    }

    for style_name, profile in style_profiles.items():
        result["style_profiles"][style_name] = {
            "paragraphs": profile["paragraphs"],
            "run_count": profile["run_count"],
            "bold_true_runs": profile["bold_true_runs"],
            "italic_true_runs": profile["italic_true_runs"],
            "fonts": dict(profile["fonts"].most_common()),
            "sizes_pt": dict(profile["sizes_pt"].most_common()),
            "alignments": dict(profile["alignments"].most_common()),
            "examples": profile["examples"],
        }

    return result


def pick_first(style_profiles: dict, keywords: list[str], fallback=None):
    for style_name, profile in sorted(
        style_profiles.items(),
        key=lambda item: item[1].get("paragraphs", 0),
        reverse=True,
    ):
        lowered = style_name.strip().lower()
        if any(keyword in lowered for keyword in keywords):
            return style_name
    return fallback


def pick_body(style_profiles: dict):
    ranked = sorted(
        style_profiles.items(),
        key=lambda item: item[1].get("paragraphs", 0),
        reverse=True,
    )
    for style_name, _profile in ranked:
        lowered = style_name.strip().lower()
        if any(bad in lowered for bad in ["heading", "title", "toc", "header", "footer", "caption"]):
            continue
        return style_name
    return ranked[0][0] if ranked else None


def run_mapping(audit_data: dict):
    style_profiles = audit_data.get("style_profiles", {})
    return {
        "template": audit_data.get("template"),
        "suggested_mapping": {
            "title": pick_first(style_profiles, ["title", "标题"]),
            "heading_1": pick_first(style_profiles, ["heading 1", "heading1", "标题1", "h1"]),
            "heading_2": pick_first(style_profiles, ["heading 2", "heading2", "标题2", "h2"]),
            "heading_3": pick_first(style_profiles, ["heading 3", "heading3", "标题3", "h3"]),
            "body": pick_body(style_profiles),
            "list": pick_first(style_profiles, ["list", "bullet", "列表", "编号"]),
            "table_text": pick_first(style_profiles, ["table", "表格", "grid"]),
            "caption": pick_first(style_profiles, ["caption", "图注", "表注"]),
            "header_footer": pick_first(style_profiles, ["header", "footer", "页眉", "页脚"]),
        },
        "notes": [
            "These mappings are heuristics derived from template style names and usage frequency.",
            "Review the suggestions before applying them to generated content.",
            "When a role is null, fall back to the closest template style manually.",
        ],
    }


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Prepare extracted content and template-format artifacts for template-driven DOCX generation."
    )
    parser.add_argument("--template", required=True, help="Path to template DOCX")
    parser.add_argument("--source", required=True, help="Path to source content document")
    parser.add_argument(
        "--workdir",
        required=True,
        help="Directory for generated preparation artifacts",
    )
    args = parser.parse_args()

    template_path = Path(args.template)
    source_path = Path(args.source)
    workdir = Path(args.workdir)

    if not template_path.exists():
        raise SystemExit(f"ERROR: template not found: {template_path}")
    if not source_path.exists():
        raise SystemExit(f"ERROR: source document not found: {source_path}")

    workdir.mkdir(parents=True, exist_ok=True)

    source_text = extract_text(source_path)
    input_md = workdir / "input.md"
    input_md.write_text(source_text, encoding="utf-8")

    template_md = workdir / "template-format.md"
    template_md.write_text(run_inspect(template_path), encoding="utf-8")

    audit_json_path = workdir / "template-format.json"
    audit_data = run_audit(template_path)
    audit_json_path.write_text(
        json.dumps(audit_data, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    style_mapping_path = workdir / "style-mapping.json"
    style_mapping = run_mapping(audit_data)
    style_mapping_path.write_text(
        json.dumps(style_mapping, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    manifest = {
        "template": str(template_path),
        "source": str(source_path),
        "artifacts": {
            "input_md": str(input_md),
            "template_format_md": str(template_md),
            "template_format_json": str(audit_json_path),
            "style_mapping_json": str(style_mapping_path),
        },
        "next_step": (
            "Use minimax-docx to apply the template, ensuring all substantive content "
            "comes from the extracted source text and all formatting follows the template."
        ),
    }
    (workdir / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )

    print(json.dumps(manifest, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
