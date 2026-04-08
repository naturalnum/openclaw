#!/usr/bin/env python3

import argparse
from copy import deepcopy
from pathlib import Path

from docx import Document


def normalize(text: str) -> str:
    return " ".join(text.split())


def paragraph_text(paragraph) -> str:
    return normalize(paragraph.text)


def find_nonempty_paragraph(document: Document, target: str):
    for paragraph in document.paragraphs:
        if paragraph_text(paragraph) == target:
            return paragraph
    return None


def clear_paragraph_keep_style(paragraph, text: str):
    sample = None
    for run in paragraph.runs:
        if run.text.strip():
            sample = run
            break

    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    new_run = paragraph.add_run(text)
    if sample is not None:
        new_run.font.name = sample.font.name
        new_run.font.size = sample.font.size
        new_run.bold = sample.bold
        new_run.italic = sample.italic


def cover_fields(source_doc: Document) -> dict[str, str]:
    values: dict[str, str] = {}
    if not source_doc.tables:
        return values
    for row in source_doc.tables[0].rows:
        cells = [normalize(cell.text) for cell in row.cells]
        if len(cells) >= 2 and cells[0]:
            values[cells[0].rstrip("：:")] = cells[1]
    return values


def find_body_child_index(document: Document, target_text: str) -> int:
    body = document.element.body
    for index, child in enumerate(body.iterchildren()):
        texts = []
        for node in child.iter():
            if node.tag.endswith("}t") and node.text:
                texts.append(node.text)
        if normalize("".join(texts)) == target_text:
            return index
    raise ValueError(f"Could not find body child for text: {target_text}")


def find_source_insert_index(document: Document, target_text: str, style_name: str | None = None, occurrence: str = "first") -> int:
    body = document.element.body
    matches = []
    for index, child in enumerate(body.iterchildren()):
        texts = []
        style_val = None
        for node in child.iter():
            if node.tag.endswith("}t") and node.text:
                texts.append(node.text)
            if node.tag.endswith("}pStyle") and style_val is None:
                style_val = node.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
        current = normalize("".join(texts))
        if current == target_text and (style_name is None or style_val == style_name):
            matches.append(index)
    if not matches:
        raise ValueError(f"Could not find source insert point: {target_text} / {style_name}")
    return matches[-1] if occurrence == "last" else matches[0]


def replace_cover(template_doc: Document, source_doc: Document):
    fields = cover_fields(source_doc)
    replacements = {
        "国家电网公司总部科技项目": "国家电网公司总部科技项目",
        "基于人机协同的配网检修运行方式智能校核与编制优化关键技术研究": fields.get("项目名称", ""),
        "《基于人机协同的配网检修运行方式智能校核与编制优化关键技术研究》总报告（1）": "《%s》技术总报告" % fields.get("项目名称", "技术报告"),
        "中国电力科学研究院有限公司": fields.get("主要承担单位", ""),
        "2026年1月": fields.get("起止时间", ""),
    }

    used_org = False
    for paragraph in template_doc.paragraphs:
        text = paragraph_text(paragraph)
        if not text:
            continue
        if text in replacements:
            new_text = replacements[text]
            if not new_text:
                continue
            if text == "中国电力科学研究院有限公司":
                if used_org:
                    continue
                used_org = True
            clear_paragraph_keep_style(paragraph, new_text)


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Use the fixed template as base and replace its body with generated technical report content."
    )
    parser.add_argument("--template", required=True)
    parser.add_argument("--source", required=True, help="Original source material DOCX")
    parser.add_argument("--content", required=True, help="Generated report content DOCX")
    parser.add_argument("--output", required=True)
    parser.add_argument(
        "--replace-from",
        default="1 前言",
        help="Paragraph text in content doc from which to begin copying into the template body.",
    )
    parser.add_argument(
        "--occurrence",
        default="last",
        choices=["first", "last"],
        help="When multiple paragraphs match --replace-from, choose which one to use.",
    )
    args = parser.parse_args()

    template_doc = Document(args.template)
    source_doc = Document(args.source)
    content_doc = Document(args.content)

    replace_cover(template_doc, source_doc)

    template_body = template_doc.element.body
    content_body = content_doc.element.body

    replace_start = find_body_child_index(template_doc, "内容摘要")
    try:
        insert_start = find_source_insert_index(content_doc, args.replace_from, occurrence=args.occurrence)
    except ValueError:
        insert_start = find_source_insert_index(content_doc, "1 项目概况")

    children = list(template_body.iterchildren())
    sect_pr = children[-1] if children and children[-1].tag.endswith("}sectPr") else None

    for child in children[replace_start:]:
        if sect_pr is not None and child is sect_pr:
            break
        template_body.remove(child)

    content_children = list(content_body.iterchildren())
    insert_nodes = []
    for child in content_children[insert_start:]:
        if child.tag.endswith("}sectPr"):
            break
        insert_nodes.append(deepcopy(child))

    if sect_pr is not None:
        for node in insert_nodes:
            template_body.insert(len(template_body) - 1, node)
    else:
        for node in insert_nodes:
            template_body.append(node)

    output = Path(args.output)
    output.parent.mkdir(parents=True, exist_ok=True)
    template_doc.save(str(output))
    print(output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
